#!/usr/bin/env python3
"""Exercise the installed V6 workflow boundary and editable-page runtime."""

from __future__ import annotations

import argparse
import json
import os
import shutil
import subprocess
from pathlib import Path

from docx import Document


def smoke(editppt: Path, output: Path) -> dict:
    if output.exists():
        shutil.rmtree(output)
    output.mkdir(parents=True)
    project = output / "project"
    page_title = "Portable clean-install workflow"
    body_text = "The V6 source, style, and editable reconstruction runtimes are available."
    source = output / "source.docx"
    document = Document()
    document.add_paragraph("\u7b2c 1 \u9875")
    document.add_paragraph(page_title)
    document.add_paragraph(body_text)
    document.save(source)
    logo = output / "logo.svg"
    logo.write_text(
        '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 120 48">'
        '<rect width="120" height="48" fill="#22577A"/></svg>',
        encoding="utf-8",
    )

    from confirm_ui.server import _wait, create_app
    from editppt.runtime.fixed_region_runtime import CONTENT_BOX, SLIDE
    from workflow_v6_source import initialize_v6_project
    from workflow_v6_state import load

    initialize_v6_project(source, logo, project)
    state = load(project)
    if len(state.get("pages", [])) != 1:
        raise RuntimeError("portable V6 initialization did not create exactly one page")

    client = create_app(project).test_client()
    recommendations = client.get("/api/recommendations").get_json()
    selected = recommendations["design_directions"]["selected"]
    candidate = recommendations["design_directions"]["candidates"][selected]
    confirmation = {
        "stage": "final",
        "direction": selected,
        "template_selection": candidate["template_selection"],
        "canvas": "ppt169",
        **{
            key: candidate[key]
            for key in (
                "visual_style", "color", "icons", "typography", "image_rendering",
                "style_axes", "layout_preferences", "information_density",
                "background_system", "image_role", "evidence_strength",
                "composition_tendency", "brand_device",
            )
        },
        "regional_style": {"enabled": False},
        "production_profile": "balanced",
        "additional_requirements": "Preserve the Word narrative and V6 fixed-layer boundary.",
    }
    response = client.post("/api/confirm", json=confirmation)
    if response.status_code != 200 or _wait(project, "final", 5) != 0:
        raise RuntimeError(f"portable V6 confirmation failed: {response.get_json()}")

    state = load(project)
    if (
        state.get("workflow_contract_version") != "word-ppt-workflow-v6"
        or state.get("style_confirmation", {}).get("status") != "confirmed"
    ):
        raise RuntimeError("portable initialization/confirmation did not reach the V6 boundary")

    page_dir = output / "editable-page"
    page_dir.mkdir()
    manifest = {
        "workflow_contract_version": "fixed-canvas-cm-v2",
        "reconstruction_contract_version": "editable-image-v3",
        "slide": dict(SLIDE),
        "content_box": dict(CONTENT_BOX),
        "source": {"width_px": 1700, "height_px": 800},
        "text_boxes": [{"object_id": "word-p1", "name": "body-paragraph-1", "text": body_text, "box_px": [80, 60, 1200, 100]}],
        "tables": [],
        "shapes": [{"object_id": "decor-1", "name": "decorative-panel", "type": "rect", "box_px": [40, 30, 1500, 650], "fill": "#F4F4F4"}],
        "images": [],
        "visual_inventory": [],
        "background_strategy": "native slide background plus editable decorative panel",
        "quality_checks": {
            "font_size_calibrated": True,
            "visual_inventory_matched": True,
            "background_strategy_checked": True,
            "shape_corner_geometry_checked": True,
        },
    }
    (page_dir / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")
    workflow_env = dict(os.environ)
    for command in ([str(editppt), "page", "build", str(page_dir)], [str(editppt), "page", "validate", str(page_dir)]):
        completed = subprocess.run(command, capture_output=True, text=True, env=workflow_env, check=False)
        if completed.returncode:
            raise RuntimeError(completed.stdout + completed.stderr)
    if not (page_dir / "page.pptx").is_file() or not (page_dir / "preview.png").is_file():
        raise RuntimeError("portable V6 object build did not create PPTX and preview")
    return {"workflow": "word-ppt-workflow-v6", "editppt": "v6-build-validate-ok"}


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--editppt", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    print(json.dumps(smoke(args.editppt.resolve(), args.output.resolve())))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
