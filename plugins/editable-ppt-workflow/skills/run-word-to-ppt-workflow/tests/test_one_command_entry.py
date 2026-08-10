"""Regression coverage for the self-contained one-command confirmation entry."""

from __future__ import annotations

import importlib.util
import json
import os
import subprocess as stdlib_subprocess
import sys
import threading
import time
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
from types import SimpleNamespace

from docx import Document
import pytest


ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
SERVER_PATH = SCRIPTS / "confirm_ui" / "server.py"
sys.path.insert(0, str(SCRIPTS))

import prepare_run  # noqa: E402
import page_requirement_summary  # noqa: E402
import natural_comment_resolver  # noqa: E402
import run_workflow  # noqa: E402
import style_recommendations  # noqa: E402
from workflow_v5_dag import DagStore  # noqa: E402
import workflow_state  # noqa: E402
from codex_subscription_runtime import CodexRuntimeUnavailable, CodexStructuredResult  # noqa: E402
from style_recommendations import build_recommendations  # noqa: E402
from test_style_contract import confirmed_result  # noqa: E402


def _source_word(path: Path) -> Path:
    document = Document()
    document.add_paragraph("第1页")
    document.add_paragraph("技术平台融资计划")
    document.add_paragraph("产品验证、专利布局和客户试点形成投资判断。")
    document.add_paragraph("第2页")
    document.add_paragraph("研发路线与实验数据")
    document.add_paragraph("展示实验数据、技术路线、产品参数和市场验证证据。")
    document.save(path)
    return path


def _logo(path: Path) -> Path:
    path.write_text(
        '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 120 48"/>',
        encoding="utf-8",
    )
    return path


def _source_word_with_ambiguous_comments(path: Path) -> Path:
    document = Document()
    for page, (title, body_text) in enumerate((
        ("技术平台融资计划", "产品验证、专利布局和客户试点形成投资判断。"),
        ("研发路线与实验数据", "展示实验数据、技术路线、产品参数和市场验证证据。"),
    ), start=1):
        document.add_paragraph(f"第{page}页")
        document.add_paragraph(title)
        body = document.add_paragraph(body_text)
        document.add_comment(body.runs, text="让这一页更有呼吸感", author="reviewer")
    document.save(path)
    return path


def _load_server():
    spec = importlib.util.spec_from_file_location("one_command_confirm_ui", SERVER_PATH)
    assert spec and spec.loader
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def test_prepare_creates_deterministic_final_recommendations_that_confirm_ui_can_read(tmp_path: Path) -> None:
    project = tmp_path / "project"
    result = prepare_run.prepare(_source_word(tmp_path / "source.docx"), project, _logo(tmp_path / "logo.svg"))

    recommendations_path = project / "confirm_ui" / "recommendations.json"
    first = recommendations_path.read_bytes()
    recommendations = json.loads(first)
    server = _load_server()
    response = server.create_app(project).test_client().get("/api/recommendations")

    assert result["next_stage"] == "style_confirmation"
    assert recommendations["stage"] == "final"
    assert len(recommendations["design_directions"]["candidates"]) >= 3
    assert server._stage2_error(recommendations) is None
    assert response.status_code == 200
    assert response.get_json()["stage"] == "final"
    assert b"created_at" not in first
    assert b"updated_at" not in first

    build_recommendations(project)

    assert recommendations_path.read_bytes() == first


def test_run_prepares_a_new_project_and_waits_at_its_single_confirmation_gate(tmp_path: Path, monkeypatch) -> None:
    project = tmp_path / "project"
    calls: list[list[str]] = []

    def start_ui(command: list[str], **_kwargs):
        calls.append(command)
        return SimpleNamespace(returncode=0, stdout="confirmation UI started", stderr="")

    monkeypatch.setattr(run_workflow.subprocess, "run", start_ui)

    result = run_workflow.run(
        _source_word(tmp_path / "source.docx"),
        _logo(tmp_path / "logo.svg"),
        project,
        no_browser=True,
    )

    assert result["stage"] == "awaiting_confirmation"
    assert (project / "confirm_ui" / "recommendations.json").is_file()
    assert calls and calls[0][1:3] == [str(run_workflow.CONFIRM_UI), "start"]


def test_production_run_imports_no_v4_runner() -> None:
    script = f"""
import importlib.abc
import sys
class RejectV4Runner(importlib.abc.MetaPathFinder):
    def find_spec(self, fullname, path=None, target=None):
        if fullname == 'production_runner':
            raise RuntimeError('V4 runner import attempted')
        return None
sys.meta_path.insert(0, RejectV4Runner())
sys.path.insert(0, {str(SCRIPTS)!r})
import run_workflow
assert not hasattr(run_workflow, 'run_production')
"""
    completed = stdlib_subprocess.run(
        [sys.executable, "-c", script], capture_output=True, text=True, check=False,
    )
    assert completed.returncode == 0, completed.stderr or completed.stdout


def test_confirmed_run_initializes_v5_and_returns_skill_ready_work(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch,
) -> None:
    word = _source_word(tmp_path / "source.docx")
    logo = _logo(tmp_path / "logo.svg")
    project = tmp_path / "project"
    prepare_run.prepare(word, project, logo)
    (project / "confirm_ui" / "result.json").write_text(
        json.dumps(confirmed_result(), ensure_ascii=False), encoding="utf-8",
    )
    monkeypatch.setattr(
        run_workflow.subprocess, "run",
        lambda *_args, **_kwargs: pytest.fail("confirmed run must not reopen UI or run V4"),
    )

    result = run_workflow.run(word, logo, project)

    assert result["stage"] == "v5_ready"
    assert result["workflow_contract_version"] == "word-ppt-workflow-v5"
    assert result["v5_state"] == "migrated"
    assert result["confirmation"] == "confirmed"
    assert result["ready_nodes"] == len(result["ready_work"]) > 0
    assert {item["kind"] for item in result["ready_work"]} == {"design"}
    assert result["dispatch_wave"]["mode"] == "parallel_pages"
    assert result["dispatch_wave"]["max_concurrency"] == 4
    assert result["dispatch_wave"]["selected_node_ids"] == [
        item["node_id"] for item in result["ready_work"]
    ]
    assert all(item["executor"] == "codex_skill_orchestrator" for item in result["ready_work"])
    assert all(item["skill"] == "generate-slide-body-image" for item in result["ready_work"])
    assert result["orchestrator_contract"] == {
        "owner": "run-word-to-ppt-workflow",
        "execution_surface": "codex_skill",
        "python_spawns_page_subagents": False,
        "reconstruction_dispatch": "one_codex_page_subagent_per_ready_page",
        "dispatch_wave_is_authoritative": True,
        "schedule_only": False,
    }
    assert (project / "04_v5" / "dag.json").is_file()


def test_repeated_confirmed_run_resumes_same_v5_dag_without_reconfirmation_or_reset(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch,
) -> None:
    word = _source_word(tmp_path / "source.docx")
    logo = _logo(tmp_path / "logo.svg")
    project = tmp_path / "project"
    prepare_run.prepare(word, project, logo)
    (project / "confirm_ui" / "result.json").write_text(
        json.dumps(confirmed_result(), ensure_ascii=False), encoding="utf-8",
    )
    monkeypatch.setattr(
        run_workflow.subprocess, "run",
        lambda *_args, **_kwargs: pytest.fail("confirmed V5 resume must not reopen UI or run V4"),
    )
    first = run_workflow.run(word, logo, project)
    store = DagStore(project)
    design = next(item for item in first["ready_work"] if item["kind"] == "design")
    store.claim(design["node_id"], worker_id="test-image2")
    store.complete(
        design["node_id"], worker_id="test-image2", result_key="sha256:" + "a" * 64,
    )
    dag_before = (project / "04_v5" / "dag.json").read_bytes()
    style_before = (project / "02_style" / "style_execution.json").read_bytes()

    second = run_workflow.run(word, logo, project)
    third = run_workflow.run(word, logo, project, execute=False)

    assert second["v5_state"] == third["v5_state"] == "resumed"
    assert (project / "04_v5" / "dag.json").read_bytes() == dag_before
    assert (project / "02_style" / "style_execution.json").read_bytes() == style_before
    node = next(
        item for item in store.snapshot()["nodes"] if item["node_id"] == design["node_id"]
    )
    assert node["status"] == "complete"
    assert node["result_key"] == "sha256:" + "a" * 64
    assert third["orchestrator_contract"]["schedule_only"] is True


def test_missing_runtime_policy_reconciles_once_then_resumes(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch,
) -> None:
    word = _source_word(tmp_path / "source.docx")
    logo = _logo(tmp_path / "logo.svg")
    project = tmp_path / "project"
    prepare_run.prepare(word, project, logo)
    (project / "confirm_ui" / "result.json").write_text(
        json.dumps(confirmed_result(), ensure_ascii=False), encoding="utf-8",
    )
    monkeypatch.setattr(
        run_workflow.subprocess, "run",
        lambda *_args, **_kwargs: pytest.fail("policy reconciliation must not reopen UI"),
    )
    first = run_workflow.run(word, logo, project)
    assert first["v5_state"] == "migrated"
    policy = project / "04_v5/runtime_policy.json"
    assert policy.is_file()
    policy.unlink()
    calls = 0
    original = run_workflow.migrate_v4_project

    def counted(path: Path):
        nonlocal calls
        calls += 1
        return original(path)

    monkeypatch.setattr(run_workflow, "migrate_v4_project", counted)
    upgraded = run_workflow.run(word, logo, project, execute=False)
    resumed = run_workflow.run(word, logo, project, execute=False)

    assert upgraded["v5_state"] == "upgraded"
    assert resumed["v5_state"] == "resumed"
    assert calls == 1
    assert json.loads(policy.read_text(encoding="utf-8")) == run_workflow._V5_POLICY


def test_run_never_treats_stderr_already_substring_as_success(tmp_path: Path, monkeypatch) -> None:
    project = tmp_path / "project"

    def failed_start(_command: list[str], **_kwargs):
        return SimpleNamespace(returncode=1, stdout="", stderr="port already belongs to a foreign process")

    monkeypatch.setattr(run_workflow.subprocess, "run", failed_start)

    with pytest.raises(RuntimeError, match="foreign process"):
        run_workflow.run(
            _source_word(tmp_path / "source.docx"),
            _logo(tmp_path / "logo.svg"),
            project,
            no_browser=True,
        )


def test_rerun_recovers_interrupted_prepare_without_rewriting_authority_or_repeating_successful_pages(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch,
) -> None:
    word = _source_word_with_ambiguous_comments(tmp_path / "source.docx")
    logo = _logo(tmp_path / "logo.svg")
    project = tmp_path / "project"
    provider_pages: list[int] = []
    page2_attempts = 0

    def interrupted_provider(_project, **kwargs):
        nonlocal page2_attempts
        prompt = kwargs["prompt"]
        page = 1 if "技术平台融资计划" in prompt else 2
        provider_pages.append(page)
        if page == 2:
            page2_attempts += 1
            if page2_attempts == 1:
                raise CodexRuntimeUnavailable("simulated App Server failure after page 1 succeeded")
        trace = {
            "runtime": "codex-app-server", "role": "comment-resolution",
            "thread_id": f"thread-{page}", "turn_id": f"turn-{page}-{page2_attempts}",
            "model": "gpt-test", "model_provider": "openai", "auth_mode": "chatgpt",
            "plan_type": "plus", "usage": {},
        }
        return CodexStructuredResult(
            value={
                "kind": "layout_override", "authority_kind": "visual_override",
                "required": True, "search_required": False, "search_query": None,
                "decisions": [{"target": "visual.layout", "action": "set", "value": "spacious"}],
            },
            thread_id=trace["thread_id"], turn_id=trace["turn_id"], model=trace["model"],
            model_provider=trace["model_provider"], auth_mode="chatgpt", plan_type="plus",
            usage={}, safe_trace=trace,
        )

    monkeypatch.setattr(natural_comment_resolver, "invoke_structured", interrupted_provider)
    with pytest.raises(natural_comment_resolver.CommentResolutionBlocked, match="simulated App Server failure"):
        run_workflow.run(word, logo, project, no_browser=True, execute=False)

    authority_paths = [
        project / "workflow_run.json",
        *sorted((project / "00_source").glob("**/*")),
        *sorted((project / "01_page_contracts").glob("**/*")),
    ]
    authority_before = {
        path.relative_to(project).as_posix(): path.read_bytes()
        for path in authority_paths if path.is_file()
    }
    assert not (project / "confirm_ui/recommendations.json").exists()

    def start_ui(command: list[str], **_kwargs):
        assert (project / "confirm_ui/recommendations.json").is_file()
        assert page_requirement_summary.verify_page_requirement_summary(
            project,
            json.loads((project / "confirm_ui/page_requirement_summary.json").read_text(encoding="utf-8")),
        )
        return SimpleNamespace(returncode=0, stdout="confirmation UI started", stderr="")

    monkeypatch.setattr(run_workflow.subprocess, "run", start_ui)
    result = run_workflow.run(word, logo, project, no_browser=True, execute=False)

    authority_after = {
        path.relative_to(project).as_posix(): path.read_bytes()
        for path in authority_paths if path.is_file()
    }
    assert result["stage"] == "awaiting_confirmation"
    assert authority_after == authority_before
    assert provider_pages == [1, 2, 2]


def test_concurrent_prepare_recovery_serializes_provider_work_and_reuses_complete_artifacts(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch,
) -> None:
    word = _source_word_with_ambiguous_comments(tmp_path / "source.docx")
    logo = _logo(tmp_path / "logo.svg")
    project = tmp_path / "project"

    monkeypatch.setattr(
        natural_comment_resolver,
        "invoke_structured",
        lambda *_args, **_kwargs: (_ for _ in ()).throw(
            CodexRuntimeUnavailable("initial prepare interruption")
        ),
    )
    with pytest.raises(natural_comment_resolver.CommentResolutionBlocked):
        run_workflow.run(word, logo, project, no_browser=True, execute=False)

    # A crashed writer and the persistent advisory-lock inode must not block recovery.
    (project / "confirm_ui").mkdir(exist_ok=True)
    (project / "confirm_ui/page_requirement_summary.json.tmp").write_text("stale", encoding="utf-8")
    (project / "confirm_ui/recommendations.json.tmp").write_text("stale", encoding="utf-8")
    (project / ".workflow_state.lock").write_bytes(b"\0")
    assert (project / ".workflow_state.lock").is_file()

    calls: list[int] = []
    calls_lock = threading.Lock()
    first_entered = threading.Event()
    release_first = threading.Event()

    def provider(_project, **kwargs):
        page = 1 if "技术平台融资计划" in kwargs["prompt"] else 2
        with calls_lock:
            calls.append(page)
            call_number = len(calls)
        if call_number == 1:
            first_entered.set()
            assert release_first.wait(10)
        trace = {
            "runtime": "codex-app-server", "role": "comment-resolution",
            "thread_id": f"thread-{page}", "turn_id": f"turn-{page}",
            "model": "gpt-test", "model_provider": "openai", "auth_mode": "chatgpt",
            "plan_type": "plus", "usage": {},
        }
        return CodexStructuredResult(
            value={
                "kind": "layout_override", "authority_kind": "visual_override",
                "required": True, "search_required": False, "search_query": None,
                "decisions": [{"target": "visual.layout", "action": "set", "value": "spacious"}],
            },
            thread_id=trace["thread_id"], turn_id=trace["turn_id"], model=trace["model"],
            model_provider=trace["model_provider"], auth_mode="chatgpt", plan_type="plus",
            usage={}, safe_trace=trace,
        )

    monkeypatch.setattr(natural_comment_resolver, "invoke_structured", provider)
    monkeypatch.setattr(
        run_workflow.subprocess,
        "run",
        lambda *_args, **_kwargs: SimpleNamespace(
            returncode=0, stdout="confirmation UI started", stderr="",
        ),
    )
    with ThreadPoolExecutor(max_workers=2) as pool:
        futures = [
            pool.submit(run_workflow.run, word, logo, project, no_browser=True, execute=False)
            for _ in range(2)
        ]
        assert first_entered.wait(10)
        time.sleep(0.2)
        with calls_lock:
            assert calls == [1]
        assert sum(future.done() for future in futures) == 0
        release_first.set()
        results = [future.result(timeout=30) for future in futures]

    artifact = json.loads(
        (project / "confirm_ui/page_requirement_summary.json").read_text(encoding="utf-8")
    )
    assert [result["stage"] for result in results] == ["awaiting_confirmation"] * 2
    assert calls == [1, 2]
    assert page_requirement_summary.verify_page_requirement_summary(project, artifact)
    assert style_recommendations.verify_prepare_artifacts(project)


def test_concurrent_first_runs_bootstrap_once_and_preserve_authority_bytes(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch,
) -> None:
    word = _source_word_with_ambiguous_comments(tmp_path / "source.docx")
    logo = _logo(tmp_path / "logo.svg")
    first_missing_parent = tmp_path / "missing-parent"
    project = first_missing_parent / "nested-parent" / "project"
    mkdir_barrier = threading.Barrier(2)
    real_mkdir = Path.mkdir

    def racing_mkdir(path: Path, *args, **kwargs):
        if path == first_missing_parent:
            mkdir_barrier.wait(timeout=10)
        return real_mkdir(path, *args, **kwargs)

    extract_calls = 0
    extract_lock = threading.Lock()
    real_extract = prepare_run.extract_auto

    def counted_extract(*args, **kwargs):
        nonlocal extract_calls
        with extract_lock:
            extract_calls += 1
        return real_extract(*args, **kwargs)

    provider_pages: list[int] = []

    def provider(_project, **kwargs):
        page = 1 if "技术平台融资计划" in kwargs["prompt"] else 2
        provider_pages.append(page)
        trace = {
            "runtime": "codex-app-server", "role": "comment-resolution",
            "thread_id": f"thread-{page}", "turn_id": f"turn-{page}",
            "model": "gpt-test", "model_provider": "openai", "auth_mode": "chatgpt",
            "plan_type": "plus", "usage": {},
        }
        return CodexStructuredResult(
            value={
                "kind": "layout_override", "authority_kind": "visual_override",
                "required": True, "search_required": False, "search_query": None,
                "decisions": [{"target": "visual.layout", "action": "set", "value": "spacious"}],
            },
            thread_id=trace["thread_id"], turn_id=trace["turn_id"], model=trace["model"],
            model_provider=trace["model_provider"], auth_mode="chatgpt", plan_type="plus",
            usage={}, safe_trace=trace,
        )

    authority_at_first_ui: dict[str, bytes] = {}
    ui_lock = threading.Lock()

    def ui_run(*_args, **_kwargs):
        with ui_lock:
            if not authority_at_first_ui:
                for relative in (
                    "workflow_run.json", "00_source/source.docx", "00_source/company_logo.svg",
                    "00_source/pages.json", "01_page_contracts/source_lock.json",
                    "01_page_contracts/page_001.json", "01_page_contracts/page_002.json",
                    "confirm_ui/page_requirement_summary.json", "confirm_ui/recommendations.json",
                ):
                    authority_at_first_ui[relative] = (project / relative).read_bytes()
        return SimpleNamespace(returncode=0, stdout="confirmation UI started", stderr="")

    monkeypatch.setattr(Path, "mkdir", racing_mkdir)
    monkeypatch.setattr(prepare_run, "extract_auto", counted_extract)
    monkeypatch.setattr(natural_comment_resolver, "invoke_structured", provider)
    monkeypatch.setattr(run_workflow.subprocess, "run", ui_run)
    with ThreadPoolExecutor(max_workers=2) as pool:
        futures = [
            pool.submit(run_workflow.run, word, logo, project, no_browser=True, execute=False)
            for _ in range(2)
        ]
        results = [future.result(timeout=30) for future in futures]

    assert extract_calls == 1
    assert provider_pages == [1, 2]
    assert [result["stage"] for result in results] == ["awaiting_confirmation"] * 2
    assert authority_at_first_ui
    assert {
        relative: (project / relative).read_bytes() for relative in authority_at_first_ui
    } == authority_at_first_ui
    assert style_recommendations.verify_prepare_artifacts(project)


@pytest.mark.parametrize("failure", ["file", "permission"])
def test_bootstrap_parent_race_does_not_hide_unsafe_mkdir_failures(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch, failure: str,
) -> None:
    raced_parent = tmp_path / "raced-parent"
    destination = raced_parent / "nested" / "project"
    real_mkdir = Path.mkdir

    def failing_mkdir(path: Path, *args, **kwargs):
        if path == raced_parent:
            if failure == "file":
                path.write_text("not a directory", encoding="utf-8")
                raise FileExistsError(str(path))
            raise PermissionError(str(path))
        return real_mkdir(path, *args, **kwargs)

    monkeypatch.setattr(Path, "mkdir", failing_mkdir)
    expected = PermissionError if failure == "permission" else ValueError
    with pytest.raises(expected):
        with workflow_state.project_bootstrap_lock(destination):
            pytest.fail("unsafe bootstrap parent must never acquire its lock")


def test_confirm_ui_link_parent_fails_closed_without_external_write(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch,
) -> None:
    word = _source_word_with_ambiguous_comments(tmp_path / "source.docx")
    logo = _logo(tmp_path / "logo.svg")
    project = tmp_path / "project"
    external = tmp_path / "external"
    external.mkdir()
    monkeypatch.setattr(
        natural_comment_resolver,
        "invoke_structured",
        lambda *_args, **_kwargs: (_ for _ in ()).throw(
            CodexRuntimeUnavailable("interrupt before confirmation artifacts")
        ),
    )
    with pytest.raises(natural_comment_resolver.CommentResolutionBlocked):
        run_workflow.run(word, logo, project, no_browser=True, execute=False)
    confirm_ui = project / "confirm_ui"
    confirm_ui.mkdir(exist_ok=True)
    for child in confirm_ui.iterdir():
        child.unlink()
    confirm_ui.rmdir()
    try:
        os.symlink(external, confirm_ui, target_is_directory=True)
    except OSError as exc:
        if os.name != "nt":
            pytest.skip(f"directory symlink unavailable: {exc}")
        created = stdlib_subprocess.run(
            ["cmd", "/c", "mklink", "/J", str(confirm_ui), str(external)],
            capture_output=True, text=True, check=False,
        )
        if created.returncode != 0:
            pytest.skip(f"directory junction unavailable: {created.stderr or created.stdout}")
    monkeypatch.setattr(natural_comment_resolver, "invoke_structured", lambda *_a, **_k: pytest.fail("unsafe path must fail before provider"))
    monkeypatch.setattr(run_workflow.subprocess, "run", lambda *_a, **_k: pytest.fail("unsafe path must fail before UI"))

    with pytest.raises(ValueError, match="link|reparse|project-local|redirect"):
        run_workflow.run(word, logo, project, no_browser=True, execute=False)
    assert list(external.iterdir()) == []


@pytest.mark.parametrize("mutation", ["false_to_zero", "true_to_one", "int_to_float"])
def test_type_confused_recommendations_fail_closed_before_ui(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch, mutation: str,
) -> None:
    word = _source_word(tmp_path / "source.docx")
    logo = _logo(tmp_path / "logo.svg")
    project = tmp_path / "project"
    if mutation == "true_to_one":
        real_recommendations = style_recommendations._recommendations

        def recommendations_with_true(contracts):
            value = real_recommendations(contracts)
            value["recommend"]["regional_style"]["enabled"] = True
            return value

        monkeypatch.setattr(style_recommendations, "_recommendations", recommendations_with_true)
    prepare_run.prepare(word, project, logo)
    path = project / "confirm_ui/recommendations.json"
    value = json.loads(path.read_text(encoding="utf-8"))
    if mutation == "false_to_zero":
        value["recommend"]["regional_style"]["enabled"] = 0
    elif mutation == "true_to_one":
        value["recommend"]["regional_style"]["enabled"] = 1
    else:
        value["design_directions"]["selected"] = float(value["design_directions"]["selected"])
    path.write_text(json.dumps(value, ensure_ascii=False), encoding="utf-8")
    monkeypatch.setattr(
        run_workflow.subprocess,
        "run",
        lambda *_args, **_kwargs: pytest.fail("type-confused recommendations must not launch UI"),
    )

    with pytest.raises(ValueError, match="existing prepare recommendations"):
        run_workflow.run(word, logo, project, no_browser=True, execute=False)


def test_rerun_freezes_browser_final_result_without_opening_ui_again(tmp_path: Path, monkeypatch) -> None:
    word = _source_word(tmp_path / "source.docx")
    logo = _logo(tmp_path / "logo.svg")
    project = tmp_path / "project"
    prepare_run.prepare(word, project, logo)
    result_path = project / "confirm_ui" / "result.json"
    result_path.write_text(json.dumps(confirmed_result(), ensure_ascii=False), encoding="utf-8")

    monkeypatch.setattr(
        run_workflow.subprocess,
        "run",
        lambda *_args, **_kwargs: pytest.fail("a finalized browser confirmation must not reopen the UI"),
    )
    result = run_workflow.run(word, logo, project, execute=False)

    state = json.loads((project / "workflow_run.json").read_text(encoding="utf-8"))
    assert state["style_confirmation"]["status"] == "confirmed"
    assert result["stage"] != "awaiting_confirmation"


def test_two_restarts_reuse_identical_frozen_style_without_starting_server(tmp_path: Path, monkeypatch) -> None:
    word = _source_word(tmp_path / "source.docx")
    logo = _logo(tmp_path / "logo.svg")
    project = tmp_path / "project"
    prepare_run.prepare(word, project, logo)
    (project / "confirm_ui" / "result.json").write_text(
        json.dumps(confirmed_result(), ensure_ascii=False), encoding="utf-8"
    )
    monkeypatch.setattr(
        run_workflow.subprocess,
        "run",
        lambda *_args, **_kwargs: pytest.fail("a frozen global confirmation must never reopen the UI"),
    )

    first = run_workflow.run(word, logo, project, execute=False)
    confirmation = (project / "02_style" / "style_confirmation.json").read_bytes()
    execution = (project / "02_style" / "style_execution.json").read_bytes()
    signature = (project / "02_style" / "style_execution.sha256").read_bytes()
    second = run_workflow.run(word, logo, project, execute=False)
    third = run_workflow.run(word, logo, project, execute=False)

    assert first["stage"] != "awaiting_confirmation"
    assert second["stage"] != "awaiting_confirmation"
    assert third["stage"] != "awaiting_confirmation"
    assert (project / "02_style" / "style_confirmation.json").read_bytes() == confirmation
    assert (project / "02_style" / "style_execution.json").read_bytes() == execution
    assert (project / "02_style" / "style_execution.sha256").read_bytes() == signature


def test_wait_ui_freezes_final_result_before_same_invocation_resumes(tmp_path: Path, monkeypatch) -> None:
    word = _source_word(tmp_path / "source.docx")
    logo = _logo(tmp_path / "logo.svg")
    project = tmp_path / "project"
    prepare_run.prepare(word, project, logo)
    commands: list[str] = []

    def confirm_ui(command: list[str], **_kwargs):
        action = command[2]
        commands.append(action)
        if action == "wait":
            (project / "confirm_ui" / "result.json").write_text(
                json.dumps(confirmed_result(), ensure_ascii=False), encoding="utf-8"
            )
        return SimpleNamespace(returncode=0, stdout="ok", stderr="")

    monkeypatch.setattr(run_workflow.subprocess, "run", confirm_ui)

    result = run_workflow.run(word, logo, project, wait_ui=True, no_browser=True, execute=False)

    state = json.loads((project / "workflow_run.json").read_text(encoding="utf-8"))
    assert commands == ["start", "wait", "shutdown"]
    assert state["style_confirmation"]["status"] == "confirmed"
    assert (project / "02_style" / "style_execution.json").is_file()
    assert result["stage"] != "await_style_confirmation"


def test_recommendations_reject_a_workflow_that_is_not_locked_for_style_confirmation(tmp_path: Path) -> None:
    project = tmp_path / "project"
    prepare_run.prepare(_source_word(tmp_path / "source.docx"), project, _logo(tmp_path / "logo.svg"))
    state_path = project / "workflow_run.json"
    state = json.loads(state_path.read_text(encoding="utf-8"))
    state["jobs"][0]["status"] = "pending"
    state_path.write_text(json.dumps(state, ensure_ascii=False), encoding="utf-8")

    with pytest.raises(ValueError, match="locked for style confirmation"):
        build_recommendations(project)


def test_recommendations_reject_a_locked_order_that_is_not_the_complete_page_sequence(tmp_path: Path) -> None:
    project = tmp_path / "project"
    prepare_run.prepare(_source_word(tmp_path / "source.docx"), project, _logo(tmp_path / "logo.svg"))
    state_path = project / "workflow_run.json"
    state = json.loads(state_path.read_text(encoding="utf-8"))
    state["pagination"]["locked_page_order"] = [2, 1]
    state_path.write_text(json.dumps(state, ensure_ascii=False), encoding="utf-8")

    with pytest.raises(ValueError, match="complete ordered page sequence"):
        build_recommendations(project)
