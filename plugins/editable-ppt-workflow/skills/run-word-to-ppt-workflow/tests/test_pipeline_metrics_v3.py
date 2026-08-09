from __future__ import annotations

import hashlib
import json
import os
import sys
import warnings
from contextlib import contextmanager
from pathlib import Path

import pytest
from docx import Document


SCRIPTS = Path(__file__).resolve().parents[1] / "scripts"
sys.path.insert(0, str(SCRIPTS))

import pipeline_metrics  # noqa: E402
from page_qa import PageQAResult, qa_issue  # noqa: E402
from prepare_run import prepare  # noqa: E402
import workflow_state  # noqa: E402
from current_contract_fixture import write_valid_generation_receipt, write_valid_qa_observation  # noqa: E402


def _read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def _report(project: Path) -> dict:
    return _read_json(project / "09_reports" / "pipeline_metrics.json")


def _page_report(project: Path, page_number: int = 1) -> dict:
    report = _report(project)
    return _read_json(project / report["page_metric_files"][str(page_number)])


def _state_revision(project: Path) -> str:
    return hashlib.sha256((project / "workflow_run.json").read_bytes()).hexdigest()


@contextmanager
def _directory_link(link: Path, target: Path):
    if os.name == "nt":
        import subprocess

        completed = subprocess.run(
            [
                "powershell.exe", "-NoProfile", "-NonInteractive", "-Command",
                "New-Item -ItemType Junction -Path $env:METRICS_LINK -Target $env:METRICS_TARGET | Out-Null",
            ],
            env={**os.environ, "METRICS_LINK": str(link), "METRICS_TARGET": str(target)},
            capture_output=True,
            text=True,
            check=False,
        )
        assert completed.returncode == 0, completed.stderr
    else:
        link.symlink_to(target, target_is_directory=True)
    try:
        yield
    finally:
        link.rmdir() if os.name == "nt" else link.unlink()


def _project(tmp_path: Path, page_count: int = 1) -> Path:
    from test_independent_page_workflow import _project as build_project

    project = build_project(tmp_path, page_count, max_concurrency=max(1, page_count))
    state_path = project / "workflow_run.json"
    state = _read_json(state_path)
    job = state["jobs"][0]
    selected_path = project / job["selected_evidence_file"]
    selected = _read_json(selected_path)
    selected.update({"available_chars": 120, "selected_chars": 45})
    selected_path.write_text(json.dumps(selected, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    job["selected_evidence_sha256"] = hashlib.sha256(selected_path.read_bytes()).hexdigest()
    state_path.write_text(json.dumps(state, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return project


def test_metrics_record_coverage_and_mapping_without_new_qa_calls() -> None:
    from pipeline_metrics import page_metrics

    value = page_metrics({
        "page_number": 2,
        "route": "hybrid",
        "coverage_result": {"passed": True, "missing": []},
        "body_image_mapping": {"mode": "contain"},
        "content_overflow": None,
    })

    assert value["coverage_passed"] is True
    assert value["coverage_missing_count"] == 0
    assert value["body_image_mapping"] == "contain"
    assert value["content_overflow"] is False


def test_metrics_use_explicit_action_counts_and_cover_current_page_state() -> None:
    advisory = qa_issue(
        code="minor_alignment",
        message="minor alignment",
        severity="advisory",
        trigger="alignment_check",
        evidence="page_002",
        confidence="high",
    )
    value = pipeline_metrics.page_metrics({
        "page_number": 2,
        "status": "content_blocked",
        "route": "hybrid",
        "attempt": 9,
        "generation": {"attempt": 8},
        "generation_calls": 2,
        "reconstruction_calls": 1,
        "semantic_calls": 7,
        "coverage_result": {"passed": False, "missing": ["date"]},
        "qa_result": {
            "status": "pass_with_advisory",
            "repair_scope": "none",
            "issues": [advisory],
        },
        "qa_path": "deterministic_only",
        "automatic_repairs_used": 1,
        "cache_hit": True,
        "body_image_mapping": {"mode": "contain"},
        "page_failure": {
            "phase": "reconstruction",
            "category": "content_overflow",
            "retryable": False,
        },
    }, {"available_chars": 120, "selected_chars": 45})

    assert value == {
        "page_number": 2,
        "status": "content_blocked",
        "route": "hybrid",
        "input_evidence_chars": 120,
        "selected_evidence_chars": 45,
        "selected_evidence_consistent": True,
        "selected_evidence_error": None,
        "generation_calls": 2,
        "reconstruction_calls": 1,
        "semantic_calls": 0,
        "image_calls": 2,
        "qa_status": "pass_with_advisory",
        "qa_path": "deterministic_only",
        "image_repairs": 1,
        "cache_hit": True,
        "coverage_passed": False,
        "coverage_missing_count": 1,
        "body_image_mapping": "contain",
        "content_overflow": True,
        "failure_phase": "reconstruction",
        "failure_category": "content_overflow",
        "failure_retryable": False,
        "warnings": [advisory],
    }


def test_each_saved_transition_rebuilds_total_and_page_metrics_from_state(tmp_path: Path) -> None:
    project = _project(tmp_path)

    action = workflow_state.next_action(project)
    first = _report(project)
    assert first["pages"][0]["status"] == "queued"
    assert first["state_revision"] == _state_revision(project)
    assert first["pages"][0]["selected_evidence_chars"] == 45
    assert first["totals"] == {
        "generation_calls": 0,
        "reconstruction_calls": 0,
        "semantic_calls": 0,
        "image_calls": 0,
        "selected_evidence_chars": 45,
        "image_repairs": 0,
    }
    assert _page_report(project) == first["pages"][0]

    generation_request = action["requests"][0]
    generation = workflow_state.dispatch(
        project, 1, "image-worker", generation_request["attempt"]
    )
    generating = _report(project)["pages"][0]
    assert generating["status"] == "generating"
    assert generating["generation_calls"] == 1
    assert generating["reconstruction_calls"] == 0
    assert _page_report(project) == generating

    image = project / "06_images" / "generated" / f"page_001_attempt_{generation['attempt']:03d}.png"
    receipt = write_valid_generation_receipt(
        project, 1, generation["attempt"], image, size=(34, 16),
    )
    workflow_state.record_generation(
        project, 1, "image-worker", generation["attempt"], image,
        generation_receipt=receipt,
    )
    qa_ready = _report(project)["pages"][0]
    assert qa_ready["status"] == "qa"
    assert qa_ready["generation_calls"] == 1

    workflow_state.record_qa(
        project,
        1,
        "image-worker",
        generation["attempt"],
        signed_invocation_bundle=write_valid_qa_observation(project, 1),
    )
    accepted = _report(project)["pages"][0]
    assert accepted["status"] == "accepted"
    assert accepted["qa_status"] == "pass"
    assert accepted["semantic_calls"] == 0

    downstream = workflow_state.next_action(project)
    assert downstream["stage"] == "reconstruction_backend_pending"
    assert downstream["requests"] == []
    still_accepted = _report(project)["pages"][0]
    assert still_accepted["status"] == "accepted"
    assert still_accepted["generation_calls"] == 1
    assert still_accepted["reconstruction_calls"] == 0
    assert _page_report(project) == still_accepted


def test_published_metric_index_and_pages_share_one_state_revision(tmp_path: Path) -> None:
    project = _project(tmp_path, 2)

    workflow_state.next_action(project)

    report = _report(project)
    revision = _state_revision(project)
    assert report["state_revision"] == revision
    assert report["snapshot"] == f"09_reports/pipeline_metrics_snapshots/{revision}"
    assert set(report["page_metric_files"]) == {"1", "2"}
    for relative in report["page_metric_files"].values():
        page = _read_json(project / relative)
        assert page["state_revision"] == revision


def test_old_reader_snapshot_remains_readable_after_new_revision_publish(tmp_path: Path) -> None:
    project = _project(tmp_path)
    action = workflow_state.next_action(project)
    old_index = _report(project)
    old_revision = old_index["state_revision"]
    old_aggregate_path = project / old_index["snapshot"] / "pipeline_metrics.json"
    old_page_path = project / old_index["page_metric_files"]["1"]

    workflow_state.dispatch(
        project, 1, "image-worker", action["requests"][0]["attempt"]
    )

    assert _report(project)["state_revision"] != old_revision
    assert _read_json(old_aggregate_path)["state_revision"] == old_revision
    assert _read_json(old_page_path)["state_revision"] == old_revision


def test_same_revision_fast_path_preserves_prior_published_snapshot(tmp_path: Path) -> None:
    project = _project(tmp_path)
    action = workflow_state.next_action(project)
    old_index = _report(project)
    old_aggregate_path = project / old_index["snapshot"] / "pipeline_metrics.json"
    old_page_path = project / old_index["page_metric_files"]["1"]

    workflow_state.dispatch(
        project, 1, "image-worker", action["requests"][0]["attempt"]
    )
    current = _report(project)
    pipeline_metrics.write_pipeline_metrics(project)

    assert _report(project) == current
    assert _read_json(old_aggregate_path)["state_revision"] == old_index["state_revision"]
    assert _read_json(old_page_path)["state_revision"] == old_index["state_revision"]


def test_current_revision_is_rebuilt_when_an_indexed_page_is_missing(tmp_path: Path) -> None:
    project = _project(tmp_path)
    workflow_state.next_action(project)
    report = _report(project)
    page_path = project / report["page_metric_files"]["1"]
    page_path.unlink()

    pipeline_metrics.write_pipeline_metrics(project)

    restored = _page_report(project)
    assert restored["state_revision"] == _state_revision(project)
    assert restored["page_number"] == 1


def test_failed_snapshot_build_keeps_previous_complete_index(tmp_path: Path, monkeypatch) -> None:
    project = _project(tmp_path, 2)
    action = workflow_state.next_action(project)
    before = _report(project)
    before_pages = {
        page: _read_json(project / relative)
        for page, relative in before["page_metric_files"].items()
    }
    real_write = pipeline_metrics.workflow_metrics._atomic_json

    def fail_second_page(project_root: Path, path: Path, value: dict) -> None:
        if path.name == "page_002.json":
            raise OSError("injected snapshot page failure")
        real_write(project_root, path, value)

    monkeypatch.setattr(pipeline_metrics.workflow_metrics, "_atomic_json", fail_second_page)
    workflow_state.dispatch(project, 1, "image-worker", action["requests"][0]["attempt"])

    assert _report(project) == before
    for page, relative in before["page_metric_files"].items():
        assert _read_json(project / relative) == before_pages[page]
    assert not (
        project / "09_reports" / "pipeline_metrics_snapshots" / _state_revision(project)
    ).exists()


def test_evidence_hash_mismatch_is_explicit_and_never_counts_mutated_text(tmp_path: Path) -> None:
    project = _project(tmp_path)
    state = workflow_state.load(project)
    selected_path = project / state["jobs"][0]["selected_evidence_file"]
    selected = _read_json(selected_path)
    selected.update({"available_chars": 9999, "selected_chars": 9999})
    selected_path.write_text(json.dumps(selected), encoding="utf-8")

    pipeline_metrics.write_pipeline_metrics(project)

    page = _report(project)["pages"][0]
    assert page["input_evidence_chars"] == 0
    assert page["selected_evidence_chars"] == 0
    assert page["selected_evidence_consistent"] is False
    assert page["selected_evidence_error"] == "sha256_mismatch"


def test_metrics_output_reparse_point_cannot_escape_project(tmp_path: Path) -> None:
    project = _project(tmp_path)
    snapshots_dir = project / "09_reports" / "pipeline_metrics_snapshots"
    snapshots_dir.parent.mkdir(parents=True)
    outside = tmp_path / "outside"
    outside.mkdir()

    with _directory_link(snapshots_dir, outside):
        workflow_state.next_action(project)

    assert list(outside.iterdir()) == []


def test_metric_write_failure_does_not_roll_back_saved_workflow_state(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch, capsys: pytest.CaptureFixture[str],
) -> None:
    project = _project(tmp_path)
    request = workflow_state.next_action(project)["requests"][0]

    def fail_metrics(*_args, **_kwargs) -> None:
        raise OSError("metrics disk unavailable")

    monkeypatch.setattr(
        pipeline_metrics.workflow_metrics, "publish_pipeline_metrics", fail_metrics
    )
    with warnings.catch_warnings():
        warnings.simplefilter("error")
        result = workflow_state.dispatch(project, 1, "image-worker", request["attempt"])

    job = workflow_state.load(project)["jobs"][0]
    assert result["state"] == "generating"
    assert job["status"] == "generating"
    assert job["assignment"] == {"agent": "image-worker", "attempt": 1, "action": "generate"}
    assert job["generation_calls"] == 1
    assert job["reconstruction_calls"] == 0
    receipt = json.loads(capsys.readouterr().err.strip().splitlines()[-1])
    assert receipt["event"] == "pipeline_metrics_refresh_failed"
    assert receipt["state_saved"] is True
    assert receipt["state_revision"] == _state_revision(project)
    assert receipt["metrics"]["error_type"] == "OSError"


def test_prepare_initializes_call_counts_and_writes_both_metric_views(tmp_path: Path) -> None:
    word = tmp_path / "source.docx"
    document = Document()
    document.add_paragraph("第1页")
    document.add_paragraph("项目进展")
    document.add_paragraph("本页正文。")
    document.save(word)
    logo = tmp_path / "logo.svg"
    logo.write_text(
        '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 120 48"/>',
        encoding="utf-8",
    )
    project = tmp_path / "prepared"

    prepare(word, project, logo)

    state = workflow_state.load(project)
    job = state["jobs"][0]
    assert job["generation_calls"] == 0
    assert job["reconstruction_calls"] == 0
    assert job["semantic_calls"] == 0
    report = _report(project)
    assert report["pages"][0]["status"] == "pending_style_confirmation"
    assert _page_report(project) == report["pages"][0]
