from __future__ import annotations

import hashlib
import json
import shutil
import sys
import time
from pathlib import Path
from types import SimpleNamespace

from docx import Document
from PIL import Image
import pytest


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

import production_runner  # noqa: E402
import workflow_state  # noqa: E402
import v4_qa_gateway  # noqa: E402
import v4_reconstruction_gateway  # noqa: E402
from prepare_run import prepare  # noqa: E402
from production_runner import run_production  # noqa: E402
from style_contract import canonical_json_bytes, compile_style_execution  # noqa: E402
from test_style_contract import confirmed_result  # noqa: E402
from cache_store import CacheStore  # noqa: E402
from current_contract_fixture import write_valid_qa_observation  # noqa: E402
from workflow_state import run_reconstruction  # noqa: E402
EDITPPT_CLI = ROOT.parent / "reconstruct-editable-slide" / "cli"
sys.path.insert(0, str(EDITPPT_CLI))
from editppt.runtime.fixed_region_runtime import CONTENT_BOX, SLIDE  # noqa: E402
from editppt.runtime.editable_page_cache import load_completed_page_package  # noqa: E402


def _raw_transport(monkeypatch, mode: str | None = None) -> None:
    def invoke(_project, *, role, prompt, **_kwargs):
        if mode == "runtime_error" and role == "qa":
            raise RuntimeError("visual backend is offline")
        if mode == "timeout" and role == "qa":
            raise TimeoutError("visual backend timed out")
        request = json.loads(prompt)
        if role == "qa":
            decision = {
                "status": "complete",
                "checks": {key: {"result": "pass", "detail": "fixture visual pass"} for key in request["check_ids"]},
                "required_image_presence": [{"asset_id": item["asset_id"], "present": True, "detail": "present"} for item in request["required_presence_images"]],
                "required_directive_results": [{"directive_id": item["directive_id"], "satisfied": True, "detail": "satisfied"} for item in request["required_directives"]],
            }
        else:
            if mode == "reconstruction_error":
                raise RuntimeError("reconstruction intentionally pending")
            text_boxes = []
            coverage = []
            for index, item in enumerate(request["authoritative_text"], start=1):
                name = f"body-text-{index:03d}"
                text_boxes.append({"object_id": item["source_id"], "name": name, "text": item["text"], "box_px": [100, 80 + 150 * (index - 1), 1500, 110], "font_size": 16, "font": "Microsoft YaHei", "color": "#111827", "bold": False, "italic": False, "align": "left", "valign": "middle", "wrap": True, "fit_text": True, "z_index": 10 + index})
                coverage.append({"source_id": item["source_id"], "text": item["text"], "object_name": name})
            decision = {"text_boxes": text_boxes, "tables": [], "shapes": [{"object_id": "visual-panel", "name": "visual-panel", "type": "rect", "box_px": [20, 20, 1860, 850], "fill": "#E8EEF4", "stroke": "#E8EEF4", "stroke_width": 0, "z_index": 0}], "images": [], "text_coverage": coverage, "table_coverage": []}
        return SimpleNamespace(value=decision, turn_id=f"turn_{role}_fixture", model="gpt-test")

    monkeypatch.setattr(v4_qa_gateway, "_invoke_structured", invoke)
    monkeypatch.setattr(v4_reconstruction_gateway, "_invoke_structured", invoke)
    monkeypatch.setattr(production_runner, "invoke_qa_gateway_worker", lambda project, work, timeout: v4_qa_gateway.invoke_builtin_gateway(project, work, timeout=timeout))
    monkeypatch.setattr(production_runner, "invoke_reconstruction_gateway_worker", lambda project, work, timeout: v4_reconstruction_gateway.invoke_builtin_gateway(project, work, timeout=timeout))


def _unavailable_qa_gateway(*_args, **_kwargs):
    raise v4_qa_gateway.GatewayUnavailable("Codex App Server unavailable")


@pytest.fixture(autouse=True)
def _clear_external_qa_credentials(monkeypatch):
    monkeypatch.delenv("OPENAI_API_KEY", raising=False)
    monkeypatch.delenv("EDITABLE_PPT_QA_PROVIDER", raising=False)
    monkeypatch.delenv("EDITABLE_PPT_QA_MODEL", raising=False)
    monkeypatch.setenv("EDITABLE_PPT_CODEX_EXECUTABLE", "missing-codex-test-executable")


def _project(tmp_path: Path, page_count: int = 1) -> Path:
    word = tmp_path / "source.docx"
    document = Document()
    for page in range(1, page_count + 1):
        document.add_paragraph(f"第{page}页")
        document.add_paragraph(f"项目进展{page}")
        document.add_paragraph(f"第{page}页首先完成方案，其次开展验证，最后完成交付。")
    document.save(word)
    logo = tmp_path / "logo.svg"
    logo.write_text('<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 120 48"/>', encoding="utf-8")
    project = tmp_path / "project"
    prepare(word, project, logo)
    return project


def _confirm(project: Path) -> None:
    execution = compile_style_execution({**confirmed_result(), "page_count": 1})
    payload = canonical_json_bytes(execution)
    digest = hashlib.sha256(payload).hexdigest()
    style = project / "02_style" / "style_execution.json"
    style.parent.mkdir(exist_ok=True)
    style.write_bytes(payload)
    (style.parent / "style_execution.sha256").write_text(digest + "\n", encoding="ascii")
    state_path = project / "workflow_run.json"
    state = json.loads(state_path.read_text(encoding="utf-8"))
    state["style_confirmation"] = {
        "status": "confirmed",
        "confirmed_at": "2026-08-01T00:00:00Z",
        "execution_file": "02_style/style_execution.json",
        "execution_sha256": digest,
    }
    state["scheduler"] = {"concurrency": 1, "configured_max": 1, "last_trigger": "test"}
    state["runtime"] = {"generation_mode": "continuous", "automatic_repair_budget": 1}
    state_path.write_text(json.dumps(state, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def _successful_backend(calls: list[list[str]]):
    def run(command, **_kwargs):
        calls.append(command)
        output = Path(command[command.index("--out") + 1])
        trace = Path(command[command.index("--trace-out") + 1])
        output.parent.mkdir(parents=True, exist_ok=True)
        generated = Image.new("RGB", (1904, 896), "#DCEAF4")
        for x in range(80, 1820, 180):
            for y in range(80, 820, 140):
                generated.paste("#23568C", (x, y, min(x + 90, 1904), min(y + 50, 896)))
        generated.save(output)
        images: list[str] = []
        roles: list[str] = []
        for index, value in enumerate(command):
            if value == "--image":
                images.append(command[index + 1])
            elif value == "--image-role":
                roles.append(command[index + 1])
        operation = command[2]
        trace.write_text(json.dumps({
            "operation": operation,
            "endpoint": "images/edits" if operation == "edit" else "images/generations",
            "model": command[command.index("--model") + 1],
            "auth": "codex_oauth",
            "input_images": [
                {
                    "role": role,
                    "path": str(Path(path).resolve()),
                    "sha256": hashlib.sha256(Path(path).read_bytes()).hexdigest(),
                }
                for path, role in zip(images, roles)
            ],
            "outputs": [{
                "path": str(output.resolve()),
                "sha256": hashlib.sha256(output.read_bytes()).hexdigest(),
            }],
        }) + "\n", encoding="utf-8")
        return SimpleNamespace(returncode=0, stderr="", stdout="")

    return run


def test_confirmed_project_runs_every_uncached_page_through_image2_then_stops_at_qa(tmp_path: Path, monkeypatch) -> None:
    project = _project(tmp_path)
    _confirm(project)
    calls: list[list[str]] = []
    monkeypatch.setattr("batch_generation.subprocess.run", _successful_backend(calls))

    result = run_production(project, finalize=False)

    assert result["stage"] == "qa_backend_pending"
    assert len(calls) == 1
    job = workflow_state.load(project)["jobs"][0]
    assert job["status"] == "qa"
    assert (project / job["generation_receipt"]["path"]).is_file()


def test_missing_codex_runtime_remains_retryable_qa_backend_pending(tmp_path: Path, monkeypatch) -> None:
    project = _project(tmp_path)
    _confirm(project)
    monkeypatch.setattr("batch_generation.subprocess.run", _successful_backend([]))
    monkeypatch.setattr(production_runner, "invoke_qa_gateway_worker", _unavailable_qa_gateway)

    result = run_production(project, finalize=False, page_timeout=2)

    assert result["stage"] == "qa_backend_pending"
    assert "Codex App Server" in result["provider_error"]
    job = workflow_state.load(project)["jobs"][0]
    assert job["status"] == "qa"
    assert job.get("automatic_repairs_used", 0) == 0


def test_small_generation_window_starts_qa_before_all_pages_finish_generating(tmp_path: Path, monkeypatch) -> None:
    project = _project(tmp_path, page_count=3)
    _confirm(project)
    calls: list[list[str]] = []
    monkeypatch.setattr("batch_generation.subprocess.run", _successful_backend(calls))

    result = run_production(project, finalize=False)

    assert result["stage"] == "page_pipeline"
    assert "Codex App Server" in result["provider_error"]
    assert len(calls) == 1
    assert [job["status"] for job in workflow_state.load(project)["jobs"]] == ["qa", "queued", "queued"]


def test_generated_prompt_requests_a_complete_body_and_excludes_fixed_layers(tmp_path: Path, monkeypatch) -> None:
    project = _project(tmp_path)
    _confirm(project)
    calls: list[list[str]] = []
    monkeypatch.setattr("batch_generation.subprocess.run", _successful_backend(calls))

    run_production(project, finalize=False)

    prompt_path = Path(calls[0][calls[0].index("--prompt-file") + 1])
    prompt = prompt_path.read_text(encoding="utf-8")
    assert "complete editable-PPT body design" in prompt
    assert "source_text is the complete factual Word authority" in prompt
    assert "page_title is drawn by the fixed title layer" in prompt
    assert "original SVG logo" in prompt
    assert "footer" in prompt and "page_number" in prompt


def test_technical_image_failure_is_paused_and_not_retried_implicitly(tmp_path: Path, monkeypatch) -> None:
    project = _project(tmp_path)
    _confirm(project)
    calls: list[list[str]] = []

    def expired(command, **_kwargs):
        calls.append(command)
        return SimpleNamespace(returncode=1, stderr="HTTP 401: token_expired", stdout="")

    monkeypatch.setattr("batch_generation.subprocess.run", expired)
    first = run_production(project, finalize=False, page_timeout=12)
    second = run_production(project, finalize=False, page_timeout=12)

    assert first["stage"] == second["stage"] == "page_blocked"
    assert first["page_failures"][0]["category"] == "authentication"
    assert len(calls) == 1


def test_accepted_page_returns_structured_reconstruction_backend_pending(tmp_path: Path, monkeypatch) -> None:
    project = _project(tmp_path)
    _confirm(project)
    calls: list[list[str]] = []
    monkeypatch.setattr("batch_generation.subprocess.run", _successful_backend(calls))
    run_production(project, finalize=False)
    run = workflow_state.load(project)
    job = run["jobs"][0]
    workflow_state.record_qa(
        project,
        1,
        "image-batch-1",
        job["attempt"],
        signed_invocation_bundle=write_valid_qa_observation(project, 1),
    )

    result = run_production(project, finalize=False)

    assert result["stage"] == "reconstruction_backend_pending"
    assert result["pending_pages"] == [1]
    assert len(result["reconstruction_work_items"]) == 1


def test_word_v4_reconstruction_rejects_unsigned_local_manifest(tmp_path: Path, monkeypatch) -> None:
    project = _project(tmp_path)
    _confirm(project)
    monkeypatch.setattr("batch_generation.subprocess.run", _successful_backend([]))
    run_production(project, finalize=False)
    job = workflow_state.load(project)["jobs"][0]
    workflow_state.record_qa(project, 1, "image-batch-1", job["attempt"], signed_invocation_bundle=write_valid_qa_observation(project, 1))
    pending = workflow_state.next_action(project)
    work_path = project / pending["reconstruction_work_items"][0]["path"]
    work = json.loads(work_path.read_text(encoding="utf-8"))
    with Image.open(project / work["accepted_body_image"]["path"]) as body:
        body_width, body_height = body.size
    text_boxes = []
    text_coverage = []
    for index, item in enumerate(work["authoritative_text"], start=1):
        name = f"body-text-{index}"
        text_boxes.append({"object_id": item["source_id"], "name": name, "text": item["text"], "font_size": 16, "box_px": [100, 70 + 120 * (index - 1), 1200, 90]})
        text_coverage.append({"source_id": item["source_id"], "text": item["text"], "object_name": name})
    manifest = {
        "artifact_version": "editable-reconstruction-manifest-v1", "work_item_sha256": hashlib.sha256(work_path.read_bytes()).hexdigest(),
        "workflow_contract_version": "fixed-canvas-cm-v2", "reconstruction_contract_version": "editable-image-v3",
        "slide": dict(SLIDE), "content_box": dict(CONTENT_BOX), "source": {"width_px": body_width, "height_px": body_height},
        "text_boxes": text_boxes, "tables": [],
        "shapes": [{"object_id": "panel-1", "name": "body-panel", "type": "rect", "box_px": [20, 20, body_width - 40, body_height - 40], "fill": "#f4f4f4", "z_index": 0}],
        "images": [], "raster_components": [], "text_coverage": text_coverage, "table_coverage": [],
    }
    manifest_path = project / "07_editable/page_001/manifest.json"
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")
    with pytest.raises(ValueError, match="signed reconstruction invocation"):
        run_reconstruction(project, 1, "reconstructor-1", manifest_path)


def _gateway_response(payload: dict) -> bytes:
    request = json.loads(payload["input"][0]["content"][0]["text"])
    decision = {
        "status": "complete",
        "checks": {key: {"result": "pass", "detail": "Checked remotely."} for key in request["check_ids"]},
        "required_image_presence": [
            {"asset_id": item["asset_id"], "present": True, "detail": "Visible."}
            for item in request["required_presence_images"]
        ],
        "required_directive_results": [
            {"directive_id": item["directive_id"], "satisfied": True, "detail": "Satisfied."}
            for item in request["required_directives"]
        ],
    }
    return json.dumps({
        "id": "resp_production_test",
        "output": [{"content": [{"type": "output_text", "text": json.dumps(decision)}]}],
    }).encode()


def test_builtin_visual_gateway_can_complete_the_sealed_qa_boundary(tmp_path: Path, monkeypatch) -> None:
    project = _project(tmp_path)
    _confirm(project)
    monkeypatch.setattr("batch_generation.subprocess.run", _successful_backend([]))
    _raw_transport(monkeypatch)
    result = run_production(project, finalize=False)

    assert result["stage"] == "assembly_pending"
    assert workflow_state.load(project)["jobs"][0]["status"] == "complete"


@pytest.mark.parametrize("failure", [RuntimeError("visual backend is offline"), TimeoutError("visual backend timed out")])
def test_visual_provider_error_stays_qa_backend_pending_without_assumed_pass(
    tmp_path: Path, monkeypatch, failure: Exception,
) -> None:
    project = _project(tmp_path)
    _confirm(project)
    monkeypatch.setattr("batch_generation.subprocess.run", _successful_backend([]))

    _raw_transport(monkeypatch, "timeout" if isinstance(failure, TimeoutError) else "runtime_error")
    result = run_production(project, finalize=False)

    assert result["stage"] == "qa_backend_pending"
    assert str(failure) in result["provider_error"]
    assert workflow_state.load(project)["jobs"][0]["status"] == "qa"
    assert workflow_state.load(project)["runtime"].get("automatic_repairs_used", 0) == 0


def test_production_runner_does_not_accept_legacy_provider_objects(tmp_path: Path, monkeypatch) -> None:
    project = _project(tmp_path)
    _confirm(project)
    monkeypatch.setattr("batch_generation.subprocess.run", _successful_backend([]))

    with pytest.raises(TypeError, match="qa_provider"):
        run_production(project, finalize=False, qa_provider=lambda *_: {"status": "complete"})


def test_hung_gateway_worker_is_terminated_without_consuming_budget(
    tmp_path: Path, monkeypatch,
) -> None:
    project = _project(tmp_path)
    _confirm(project)
    monkeypatch.setattr("batch_generation.subprocess.run", _successful_backend([]))
    hanging_worker = tmp_path / "hang_gateway.py"
    hanging_worker.write_text("import time\ntime.sleep(60)\n", encoding="utf-8")
    monkeypatch.setattr(v4_qa_gateway, "__file__", str(hanging_worker))

    started = time.monotonic()
    result = run_production(project, finalize=False, page_timeout=0.25)
    elapsed = time.monotonic() - started

    assert result["stage"] == "qa_backend_pending"
    assert "timeout" in result["provider_error"].lower()
    assert elapsed < 10.0
    job = workflow_state.load(project)["jobs"][0]
    assert job["status"] == "qa"
    assert job.get("automatic_repairs_used", 0) == 0
    assert "qa_receipt" not in job


def test_production_module_has_no_legacy_native_or_overlay_backend() -> None:
    assert not hasattr(production_runner, "build_native_page")
    assert not hasattr(production_runner, "build_overlay_page")
    assert not hasattr(production_runner, "choose_page_route")


def test_unconfirmed_project_remains_at_the_style_gate(tmp_path: Path) -> None:
    project = _project(tmp_path)

    result = run_production(project, finalize=False)

    assert result["stage"] == "await_style_confirmation"
    assert result["project"] == str(project.resolve())


def test_seeded_final_pptx_cannot_escape_pending_v4_backends(tmp_path: Path) -> None:
    project = _project(tmp_path)
    _confirm(project)
    seeded = project / "08_final" / "seeded.pptx"
    seeded.parent.mkdir(parents=True, exist_ok=True)
    seeded.write_bytes(b"manually seeded legacy deck")
    state_path = project / "workflow_run.json"
    state = json.loads(state_path.read_text(encoding="utf-8"))
    state["final_pptx"] = seeded.relative_to(project).as_posix()
    state_path.write_text(json.dumps(state, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

    result = run_production(project)

    assert result["stage"] == "authority_changed"
    assert result["reason"] == "completed_project_authorities_changed"


def test_full_page_cache_hit_cannot_trigger_finalization_while_backend_is_pending(
    tmp_path: Path, monkeypatch,
) -> None:
    project = _project(tmp_path)
    _confirm(project)
    workflow_state.next_action(project)
    state = workflow_state.load(project)
    job = state["jobs"][0]
    store = CacheStore(project)
    with store.staging("pages", job["cache"]["key"]) as staged:
        artifact = staged / "legacy-page.pptx"
        artifact.write_bytes(b"legacy full-page cache")
        store.seal("pages", job["cache"]["key"], staged, {
            "schema_version": 1,
            "cache_identity": job["cache"]["identity"],
            "outputs": {"reconstruction": artifact.name},
            "files": [{
                "path": artifact.name,
                "sha256": hashlib.sha256(artifact.read_bytes()).hexdigest(),
            }],
        })
    assert hasattr(production_runner, "finalize_project")
    monkeypatch.setattr("batch_generation.subprocess.run", _successful_backend([]))
    monkeypatch.setattr(production_runner, "invoke_qa_gateway_worker", _unavailable_qa_gateway)

    result = run_production(project, finalize=True)

    assert result["stage"] == "qa_backend_pending"
    assert "Codex App Server" in result["provider_error"]
    assert workflow_state.load(project)["jobs"][0]["status"] != "complete"
