from __future__ import annotations

import inspect
import json
import math
import sys
from pathlib import Path

import pytest

pytestmark = pytest.mark.skip(reason="retired V4 production workflow; V6 has independent acceptance coverage")


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

import run_workflow  # noqa: E402
import v4_reconstruction_gateway  # noqa: E402
import workflow_state  # noqa: E402
from prepare_run import prepare  # noqa: E402
from production_runner import run_production  # noqa: E402
from test_v4_one_command_orchestration import (  # noqa: E402
    _accepted_project, _authorize_gateway, _provider_response,
)


def test_word_v4_reconstruction_api_and_cli_accept_only_signed_gateway_bundle() -> None:
    signature = inspect.signature(workflow_state.run_reconstruction)
    assert "gateway_bundle" in signature.parameters
    assert "manifest" not in signature.parameters
    source = inspect.getsource(workflow_state)
    assert 'reconstruct.add_argument("--gateway-bundle"' in source
    assert 'reconstruct.add_argument("--manifest"' not in source


def test_completed_fast_path_revalidates_changed_style_before_returning_complete(tmp_path: Path, monkeypatch) -> None:
    project = _accepted_project(tmp_path, monkeypatch)
    _authorize_gateway(monkeypatch)
    assert run_production(project, finalize=True)["stage"] == "complete"
    state = workflow_state.load(project)
    style = project / state["style_confirmation"]["execution_file"]
    value = json.loads(style.read_text(encoding="utf-8"))
    value["hard_constraints"]["title_color"] = "#654321"
    style.write_text(json.dumps(value, ensure_ascii=False), encoding="utf-8")
    changed = style.read_bytes()

    with pytest.raises(ValueError, match="style"):
        run_production(project, finalize=True)
    assert style.read_bytes() == changed


def test_existing_project_rejects_different_command_word_or_logo_without_ignoring_new_authority(tmp_path: Path) -> None:
    original_word = tmp_path / "original.docx"
    original_logo = tmp_path / "original.svg"
    from docx import Document
    document = Document(); document.add_paragraph("第1页"); document.add_paragraph("原始标题"); document.add_paragraph("原始正文"); document.save(original_word)
    original_logo.write_text('<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 10 10"/>', encoding="utf-8")
    project = tmp_path / "project"
    prepare(original_word, project, original_logo)
    changed_word = tmp_path / "changed.docx"
    changed = Document(); changed.add_paragraph("第1页"); changed.add_paragraph("新标题"); changed.add_paragraph("新正文"); changed.save(changed_word)
    changed_logo = tmp_path / "changed.svg"
    changed_logo.write_text('<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 20 10"/>', encoding="utf-8")

    word_result = run_workflow.run(changed_word, original_logo, project, no_browser=True, execute=False)
    logo_result = run_workflow.run(original_word, changed_logo, project, no_browser=True, execute=False)

    assert word_result["stage"] == logo_result["stage"] == "authority_changed"
    assert changed_word.read_bytes() != (project / "00_source/source.docx").read_bytes()
    assert changed_logo.read_bytes() != (project / "00_source/company_logo.svg").read_bytes()


@pytest.mark.parametrize("mutation", ["extra", "missing", "nan", "bad_box", "duplicate"])
def test_provider_manifest_is_closed_and_semantically_validated_before_materialize(tmp_path: Path, monkeypatch, mutation: str) -> None:
    project = _accepted_project(tmp_path, monkeypatch)
    action = workflow_state.next_action(project)
    work_path = project / action["reconstruction_work_items"][0]["path"]
    work = json.loads(work_path.read_text(encoding="utf-8"))
    request = v4_reconstruction_gateway._provider_request(project, work, "gpt-4.1")
    decision = json.loads(json.loads(_provider_response(request))["output"][0]["content"][0]["text"])
    if mutation == "extra": decision["text_boxes"][0]["unexpected"] = True
    elif mutation == "missing": decision["text_boxes"][0].pop("font")
    elif mutation == "nan": decision["text_boxes"][0]["font_size"] = math.nan
    elif mutation == "bad_box": decision["text_boxes"][0]["box_px"] = [-1, 0, 99999, 1]
    else: decision["text_boxes"].append(dict(decision["text_boxes"][0]))

    with pytest.raises(ValueError, match="provider manifest"):
        v4_reconstruction_gateway.validate_provider_manifest(decision, source=request["input"][0]["content"][0]["text"])


def test_reconstruction_schema_is_app_server_compatible_and_local_validation_preserves_box_semantics(
    tmp_path: Path, monkeypatch,
) -> None:
    schema = v4_reconstruction_gateway._manifest_schema()
    box_schema = schema["properties"]["text_boxes"]["items"]["properties"]["box_px"]
    assert isinstance(box_schema["items"], dict)
    assert "prefixItems" not in box_schema
    assert "uniqueItems" not in json.dumps(schema)

    project = _accepted_project(tmp_path, monkeypatch)
    action = workflow_state.next_action(project)
    work_path = project / action["reconstruction_work_items"][0]["path"]
    work = json.loads(work_path.read_text(encoding="utf-8"))
    request = v4_reconstruction_gateway._provider_request(project, work, "gpt-test")
    decision = json.loads(json.loads(_provider_response(request))["output"][0]["content"][0]["text"])
    decision["text_boxes"][0]["box_px"] = [0, 0, 0, 1]

    with pytest.raises(ValueError, match="positive width and height"):
        v4_reconstruction_gateway.validate_provider_manifest(
            decision, source=request["input"][0]["content"][0]["text"]
        )

    valid_line = json.loads(json.loads(_provider_response(request))["output"][0]["content"][0]["text"])
    valid_line["shapes"].append({
        "object_id": "divider-1", "name": "horizontal-divider", "type": "line",
        "box_px": [0, 120, 600, 0], "fill": "#000000", "stroke": "#000000",
        "stroke_width": 1, "z_index": 5,
    })
    v4_reconstruction_gateway.validate_provider_manifest(
        valid_line, source=request["input"][0]["content"][0]["text"]
    )

    duplicate_coverage = json.loads(json.loads(_provider_response(request))["output"][0]["content"][0]["text"])
    duplicate_coverage["text_coverage"].append(dict(duplicate_coverage["text_coverage"][0]))
    with pytest.raises(ValueError, match="duplicate coverage"):
        v4_reconstruction_gateway.validate_provider_manifest(
            duplicate_coverage, source=request["input"][0]["content"][0]["text"]
        )

    extra_text = json.loads(json.loads(_provider_response(request))["output"][0]["content"][0]["text"])
    extra_text["text_boxes"].append({**extra_text["text_boxes"][0], "object_id": "decorative-label", "name": "plus-sign", "text": "+"})
    normalized = v4_reconstruction_gateway._drop_uncovered_factual_objects(extra_text)
    assert all(item["name"] != "plus-sign" for item in normalized["text_boxes"])
    v4_reconstruction_gateway._strict_coverage_maps(work, normalized)


def test_active_image_reconstruction_skill_is_explicitly_v4_safe() -> None:
    skill = (ROOT.parent / "reconstruct-editable-slide/SKILL.md").read_text(encoding="utf-8")
    assert "word-ppt-workflow-v4" in skill
    assert "relative aspect error" in skill
    assert "repair or block" in skill
    assert "standalone" in skill and "contain" in skill
