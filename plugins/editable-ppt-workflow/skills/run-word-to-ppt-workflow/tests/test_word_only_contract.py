"""Regression tests for the current-only paginated Word prepare contract."""

from __future__ import annotations

import inspect
import json
import subprocess
import sys
from pathlib import Path

import pytest
from docx import Document
from jsonschema import Draft202012Validator
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
sys.path.insert(0, str(SCRIPTS))

import prepare_run  # noqa: E402
import build_page_contracts as page_contract_builder  # noqa: E402
from build_page_contracts import build as build_page_contracts  # noqa: E402
from build_page_contracts import split_page_title_body  # noqa: E402
from extract_docx_pages import extract_auto  # noqa: E402
from page_assets import classify_page_asset  # noqa: E402


PAGE_AUTHORITY_FIXTURES = json.loads(
    (Path(__file__).parent / "fixtures" / "page_authority_regressions.json").read_text(encoding="utf-8")
)


def test_single_long_marker_page_derives_title_and_preserves_full_body():
    source = "聚焦港澳、内地、国际三地市场，以承建地产为主业，培育实业、金融投资业务，自然增长、联营合作增长、并购增长三种增长模式并举，实现利润10%增长，成为建筑和地产行业最具影响力和竞争力的国际化公司。"
    title, body = split_page_title_body(source, 2, pagination_mode="ordered_markers")
    assert title == "聚焦三地市场与三种增长模式"
    assert body == source


@pytest.mark.parametrize(
    "case",
    PAGE_AUTHORITY_FIXTURES["split_cases"],
    ids=lambda case: case["id"],
)
def test_page_title_split_preserves_word_authority(case: dict) -> None:
    """Dropping a non-heading source block or changing established pagination behavior breaks authority."""
    title, body = split_page_title_body(
        case["source"],
        case["page_number"],
        pagination_mode=case["pagination_mode"],
        page_comments=case.get("page_comments"),
    )

    assert title == case["expected_title"]
    assert body == case["expected_body"]


@pytest.mark.parametrize(
    ("case_id", "expected_origin"),
    [
        ("single_long_marker_page", "derived_from_body"),
        ("generic_ordered_marker_heading", "explicit_word_heading"),
        ("short_explicit_heading", "explicit_word_heading"),
        ("comment_title_override", "comment_override"),
        ("single_line_image_caption", "derived_from_body"),
    ],
)
def test_page_contract_stores_closed_title_origin(
    tmp_path: Path, case_id: str, expected_origin: str
) -> None:
    case = next(item for item in PAGE_AUTHORITY_FIXTURES["split_cases"] if item["id"] == case_id)
    source = tmp_path / f"{case_id}.json"
    source.write_text(
        json.dumps(
            {
                "pagination_mode": case["pagination_mode"],
                "pages": [
                    {
                        "page_number": 1,
                        "blocks": [{"type": "paragraph", "text": case["source"]}],
                        "page_comments": case.get("page_comments", []),
                    }
                ],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    output = tmp_path / f"{case_id}-contracts"
    assert build_page_contracts(source, output) == 1
    contract = json.loads((output / "page_001.json").read_text(encoding="utf-8"))
    assert contract["title_origin"] == expected_origin
    assert contract["page_title"] == case["expected_title"]
    assert contract["body_text"] == case["expected_body"]
    assert contract["source_text"] == case["source"]


def test_source_body_coverage_rejects_a_truncated_explicit_heading_projection() -> None:
    with pytest.raises(ValueError, match="loses non-heading Word text"):
        page_contract_builder.validate_source_body_coverage(
            "第一页标题\n第一段正文。\n第二段正文。",
            "第一页标题",
            "第一段正文。",
            "explicit_word_heading",
            page_number=1,
        )


def test_title_only_ordered_marker_page_builds_with_empty_body_and_full_source(tmp_path: Path) -> None:
    source = tmp_path / "title-only.json"
    source.write_text(
        json.dumps(
            {
                "pagination_mode": "ordered_markers",
                "pages": [
                    {
                        "page_number": 1,
                        "blocks": [{"type": "paragraph", "text": "年度经营计划"}],
                    }
                ],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    output = tmp_path / "title-only-contracts"
    assert build_page_contracts(source, output) == 1
    contract = json.loads((output / "page_001.json").read_text(encoding="utf-8"))
    assert contract["page_title"] == "年度经营计划"
    assert contract["title_origin"] == "explicit_word_heading"
    assert contract["body_text"] == ""
    assert contract["source_text"] == "年度经营计划"


def make_marked_word(path: Path) -> Path:
    document = Document()
    document.add_paragraph("第1页")
    document.add_paragraph("第一页内容必须留在第一页。")
    document.add_paragraph("第2页")
    document.add_paragraph("第二页内容必须留在第二页。")
    document.save(path)
    return path


def make_logo(path: Path) -> Path:
    path.write_text('<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 100 40"><rect width="100" height="40"/></svg>', encoding="utf-8")
    return path


def test_prepare_requires_an_svg_logo(tmp_path: Path) -> None:
    with pytest.raises(ValueError, match="Logo is required"):
        prepare_run.prepare(make_marked_word(tmp_path / "source.docx"), tmp_path / "project")


def test_prepare_accepts_word_output_and_required_svg_logo_and_locks_ordered_markers(tmp_path: Path):
    """Removing required inputs or reordering marker pages is a contract break."""
    word = make_marked_word(tmp_path / "source.docx")
    project = tmp_path / "project"

    assert list(inspect.signature(prepare_run.prepare).parameters) == ["word", "output", "logo"]
    result = prepare_run.prepare(word, project, make_logo(tmp_path / "logo.svg"))

    pages = json.loads((project / "00_source" / "pages.json").read_text(encoding="utf-8"))
    assert result["page_count"] == 2
    assert pages["pagination_mode"] == "explicit_text_markers"
    assert [page["page_number"] for page in pages["pages"]] == [1, 2]
    assert pages["pages"][0]["blocks"][0]["text"] == "第一页内容必须留在第一页。"
    assert pages["pages"][1]["blocks"][0]["text"] == "第二页内容必须留在第二页。"


def test_prepare_cli_creates_a_project_without_legacy_inputs(tmp_path: Path):
    """Requiring or silently accepting a legacy input breaks the public CLI contract."""
    word = make_marked_word(tmp_path / "source.docx")
    project = tmp_path / "project"
    command = [
        sys.executable,
        str(SCRIPTS / "word_to_editable_ppt.py"),
        "v6",
        "init",
        "--word",
        str(word),
        "--project",
        str(project),
        "--logo",
        str(make_logo(tmp_path / "logo.svg")),
    ]

    completed = subprocess.run(command, capture_output=True, text=True, check=False)

    assert completed.returncode == 0, completed.stderr
    assert len(json.loads(completed.stdout)["pages"]) == 2
    assert (project / "workflow_v6.json").is_file()

    rejected = subprocess.run(
        [*command, "--style-reference", "unused.png", "--company-logo", "unused.png"],
        capture_output=True,
        text=True,
        check=False,
    )
    assert rejected.returncode != 0
    assert "unrecognized arguments" in rejected.stderr


def test_workflow_cli_waits_for_style_confirmation_and_rejects_legacy_commands(tmp_path: Path):
    """A fresh project must not enter a removed master, sample, or global-QA workflow."""
    project = tmp_path / "project"
    init = [
        sys.executable, str(SCRIPTS / "word_to_editable_ppt.py"), "v6", "init",
        "--word", str(make_marked_word(tmp_path / "source.docx")),
        "--logo", str(make_logo(tmp_path / "logo.svg")),
        "--project", str(project),
    ]
    assert subprocess.run(init, capture_output=True, text=True, check=False).returncode == 0
    completed = subprocess.run(
        [sys.executable, str(SCRIPTS / "word_to_editable_ppt.py"), "v6", "status", "--project", str(project)],
        capture_output=True, text=True, check=False,
    )
    assert completed.returncode == 0, completed.stderr
    assert json.loads(completed.stdout)["next_action"] == "confirm_global_style"

    for legacy_command in ["confirm-master", "record-sample", "record-global-qa"]:
        rejected = subprocess.run(
            [sys.executable, str(SCRIPTS / "word_to_editable_ppt.py"), legacy_command, "--project", str(project)],
            capture_output=True,
            text=True,
            check=False,
        )
        assert rejected.returncode != 0
        assert "invalid choice" in rejected.stderr


def test_physical_fallback_is_used_only_when_markers_are_absent(tmp_path: Path, monkeypatch: pytest.MonkeyPatch):
    """Invalid markers must remain an error; only an unmarked document may use physical pagination."""
    unmarked = Document()
    unmarked.add_paragraph("没有页码标记的正文")
    unmarked_path = tmp_path / "unmarked.docx"
    unmarked.save(unmarked_path)
    physical_pages = {
        "schema_version": "1.1",
        "source_file": unmarked_path.name,
        "pagination_mode": "physical_rendered_pages",
        "pagination_backend": "test-renderer",
        "page_count": 1,
        "pages": [{"page_number": 1, "blocks": [{"type": "paragraph", "text": "物理分页内容"}], "must_keep": [], "page_purpose": "待人工填写"}],
    }
    monkeypatch.setattr("extract_docx_pages.extract_physical", lambda _path: physical_pages)

    assert extract_auto(unmarked_path) == physical_pages

    invalid = Document()
    invalid.add_paragraph("第2页")
    invalid.add_paragraph("顺序错误")
    invalid_path = tmp_path / "invalid.docx"
    invalid.save(invalid_path)
    with pytest.raises(ValueError, match="ordered and consecutive"):
        extract_auto(invalid_path)


def test_prepared_state_has_the_word_only_shape(tmp_path: Path):
    """Reintroducing old master, sample, logo, or visual-DNA state is a contract break."""
    project = tmp_path / "project"
    prepare_run.prepare(make_marked_word(tmp_path / "source.docx"), project, make_logo(tmp_path / "logo.svg"))

    state = json.loads((project / "workflow_run.json").read_text(encoding="utf-8"))
    assert {"workflow_contract_version", "word_source", "pagination", "style_confirmation", "jobs", "final_pptx"} <= state.keys()
    assert not {"master_jobs", "sample_status", "company_logo", "visual_dna_receipt"} & state.keys()
    assert state["word_source"]["path"] == "00_source/source.docx"
    assert state["pagination"]["locked_page_order"] == [1, 2]
    assert state["style_confirmation"]["status"] == "pending"
    assert [job["page_number"] for job in state["jobs"]] == [1, 2]
    assert [job["expected_output"] for job in state["jobs"]] == [
        "06_images/generated/page_001.png",
        "06_images/generated/page_002.png",
    ]
    assert state["final_pptx"] is None

    template = ROOT / "template"
    assert (template / "06_images" / "generated").is_dir()
    assert (template / "08_final").is_dir()
    assert not (template / "06_images" / "approved").exists()
    assert not (template / "06_images" / "draft").exists()
    assert not (template / "09_deliverables").exists()


def test_inline_word_image_is_extracted_and_bound_only_to_its_locked_page(tmp_path: Path):
    """A Word image must never leak into another independently generated slide."""
    image_path = tmp_path / "chart.png"
    Image.new("RGB", (80, 40), "#22577A").save(image_path)
    document = Document()
    document.add_paragraph("第1页")
    document.add_paragraph("根据下图说明第一阶段进展。")
    document.add_picture(str(image_path))
    document.add_paragraph("第2页")
    document.add_paragraph("第二页不应收到第一页图片。")
    word = tmp_path / "with-image.docx"
    document.save(word)

    project = tmp_path / "project"
    prepare_run.prepare(word, project, make_logo(tmp_path / "logo.svg"))
    first = json.loads((project / "01_page_contracts/page_001.json").read_text(encoding="utf-8"))
    second = json.loads((project / "01_page_contracts/page_002.json").read_text(encoding="utf-8"))

    assert len(first["asset_bindings"]) == 1
    binding = first["asset_bindings"][0]
    assert binding["asset_role"] == "mandatory_inline_image"
    assert binding["processing"] == "direct_image"
    assert binding["use_policy"] == "required"
    assert binding["advisories"] == []
    assert (project / binding["generation_input"]["relative_path"]).read_bytes() == image_path.read_bytes()
    assert second["asset_bindings"] == []
    manifest = json.loads((project / "00_source/source_asset_manifest.json").read_text(encoding="utf-8"))
    schema = json.loads((ROOT / "schemas/source_asset_manifest.schema.json").read_text(encoding="utf-8"))
    assert not list(Draft202012Validator(schema).iter_errors(manifest))
    page_schema = json.loads((ROOT / "schemas/page_contract.schema.json").read_text(encoding="utf-8"))
    assert not list(Draft202012Validator(page_schema).iter_errors(first))


def test_word_comment_is_page_local_instruction_and_never_body_text(tmp_path: Path):
    document = Document()
    document.add_paragraph("第1页")
    document.add_paragraph("第一页标题")
    body = document.add_paragraph("正文内容。")
    document.add_comment(body.runs, text="本页采用时间轴，但不要改变主结论。", author="reviewer")
    document.add_paragraph("第2页")
    document.add_paragraph("第二页标题")
    document.add_paragraph("第二页正文。")
    word = tmp_path / "with-comments.docx"
    document.save(word)

    project = tmp_path / "project"
    prepare_run.prepare(word, project, make_logo(tmp_path / "logo.svg"))
    first = json.loads((project / "01_page_contracts/page_001.json").read_text(encoding="utf-8"))
    second = json.loads((project / "01_page_contracts/page_002.json").read_text(encoding="utf-8"))

    assert [item["text"] for item in first["page_comments"]] == ["本页采用时间轴，但不要改变主结论。"]
    assert "时间轴" not in first["source_text"]
    assert second["page_comments"] == []
    assert first["requirement_precedence"] == [
        "fixed_hard_rules", "ui_global_soft_preferences", "model_creative_freedom", "word_page_comments"
    ]
    schema = json.loads((ROOT / "schemas/page_contract.schema.json").read_text(encoding="utf-8"))
    assert not list(Draft202012Validator(schema).iter_errors(first))
    assert not list(Draft202012Validator(schema).iter_errors(second))


def test_word_table_literal_pipe_is_escaped_losslessly_in_the_page_contract(tmp_path: Path):
    """A literal pipe in a Word cell must remain cell text, not become a false column boundary."""
    document = Document()
    document.add_paragraph("第1页")
    document.add_paragraph("Revenue summary")
    document.add_paragraph("Revenue was 100.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Revenue | recurring"
    table.cell(1, 1).text = "100"
    word = tmp_path / "literal-pipe.docx"
    document.save(word)
    project = tmp_path / "project"

    prepare_run.prepare(word, project, make_logo(tmp_path / "logo.svg"))

    contract = json.loads((project / "01_page_contracts/page_001.json").read_text(encoding="utf-8"))
    assert contract["source_tables"] == [
        "| Metric | Value |\n| --- | --- |\n| Revenue \\| recurring | 100 |"
    ]


@pytest.mark.parametrize(
    ("media_type", "expected_role", "expected_processing"),
    [
        ("application/pdf", "document_source", "extract_content"),
        ("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "data_source", "extract_content"),
        ("application/vnd.openxmlformats-officedocument.presentationml.presentation", "document_source", "extract_content"),
        ("application/octet-stream", "unsupported", "unavailable"),
    ],
)
def test_page_attachment_classification_is_explicit_and_non_blocking(
    media_type: str, expected_role: str, expected_processing: str
):
    """Unreadable attachments should be visible to QA without stopping other pages."""
    result = classify_page_asset(media_type, binding_status="bound", has_generation_input=False)
    assert result["asset_role"] == expected_role
    assert result["processing"] == expected_processing
    assert result["blocking"] is False
    assert bool(result["advisories"]) is (expected_processing == "unavailable")
