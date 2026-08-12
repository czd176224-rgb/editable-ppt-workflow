from __future__ import annotations

import hashlib
import json
import sys
from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
if str(SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCRIPTS))

from workflow_v6_source import compile_effective_page, initialize_v6_project  # noqa: E402
import workflow_v6_source  # noqa: E402


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    relationship_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    value = OxmlElement("w:t")
    value.text = text
    run.append(value)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def test_comments_override_word_and_unavailable_attachment_invalidates_only_reference():
    value = compile_effective_page(
        page_number=1,
        word_text="原文事实为甲。",
        comments=[{"comment_id": "7", "text": "将事实改为乙，并引用附件。"}],
        references=[{"kind": "attachment", "status": "unavailable"}],
        attachment_links=[],
    )
    assert value["comment_directives"][0]["precedence"] == "overrides_word_content"
    assert value["invalidated_requirements"] == [{
        "comment_id": "7",
        "kind": "attachment_reference",
        "reason": "attachment_unavailable",
    }]


def test_search_request_records_only_page_and_purpose():
    value = compile_effective_page(
        page_number=3,
        word_text="正文",
        comments=[{"comment_id": "1", "text": "搜索相关新闻图片作为参考。"}],
        references=[],
        attachment_links=[],
    )
    assert value["search_requests"] == [
        {"page_number": 3, "purpose": "搜索相关新闻图片作为参考。"}
    ]


def test_initialize_v6_project_uses_explicit_word_pages_without_legacy_state(tmp_path: Path):
    word = tmp_path / "input.docx"
    logo = tmp_path / "logo.svg"
    project = tmp_path / "project"
    document = Document()
    document.add_paragraph("第1页")
    document.add_paragraph("第一页标题")
    document.add_paragraph("第一页正文")
    document.add_paragraph("第2页 PPT")
    document.add_paragraph("第二页标题")
    document.add_paragraph("第二页正文")
    document.save(word)
    logo.write_text(
        '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 100 20"></svg>',
        encoding="utf-8",
    )

    state = initialize_v6_project(word, logo, project)

    assert state["workflow_contract_version"] == "word-ppt-workflow-v6"
    assert [page["title"] for page in state["pages"]] == ["第一页标题", "第二页标题"]
    assert not (project / "workflow_run.json").exists()
    page = json.loads(
        (project / "02_v6" / "page_sources" / "page_001.json").read_text(encoding="utf-8")
    )
    assert "第一页正文" in page["word_original"]
    assert page["fixed_page_title"] == "第一页标题"
    assert page["body_render_content"] == "第一页正文"
    materials = json.loads(
        (project / "02_v6" / "page_materials" / "page_001.json").read_text(
            encoding="utf-8"
        )
    )
    assert materials["fixed_page_title"] == "第一页标题"
    assert materials["effective_body"] == "第一页正文"
    assert materials["reference_images"] == []
    assert state["page_materials_status"] == "pre_confirmation"


def test_initialize_v6_project_compiles_comment_resolution_into_confirmed_materials(tmp_path: Path):
    """Initialization must carry concrete pre-UI inputs without leaking reviewer prose to materials."""
    word = tmp_path / "input.docx"
    logo = tmp_path / "logo.svg"
    project = tmp_path / "project"
    document = Document()
    document.add_paragraph("第1页")
    title = document.add_paragraph("Growth")
    body = document.add_paragraph("Context remains. Revenue was 20%.")
    closing = document.add_paragraph("Closing remains.")
    link = document.add_paragraph()
    _add_hyperlink(link, "https://example.test/report-b", "report-b")
    document.add_comment(
        [body.runs[0]],
        "Change the revenue fact Revenue was 20% to Revenue was 30%.",
        author="Reviewer",
        initials="RV",
    )
    document.add_comment(
        [closing.runs[0]],
        "Use attachment attachment-01 rows 2, 4 fields Revenue, Margin.",
        author="Reviewer",
        initials="RV",
    )
    document.add_comment(
        [title.runs[0]],
        "[search-evidence:growth evidence]",
        author="Reviewer",
        initials="RV",
    )
    document.save(word)
    logo.write_text('<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 100 20"/>', encoding="utf-8")

    initialize_v6_project(word, logo, project)

    materials = json.loads(
        (project / "02_v6/page_materials/page_001.json").read_text(encoding="utf-8")
    )
    expected_id = "search-request-" + hashlib.sha256(b"growth evidence").hexdigest()[:16]
    assert materials["effective_body"] == "Context remains. Revenue was 30%.\n\nClosing remains.\n\nreport-b"
    assert {item["attachment_id"] for item in materials["attachment_extracts"] if "attachment_id" in item} == {"attachment-01"}
    request = next(item for item in materials["attachment_extracts"] if item.get("attachment_id") == "attachment-01")
    assert request["selector"] == "selected_rows"
    assert request["rows"] == [2, 4]
    assert request["fields"] == ["Revenue", "Margin"]
    assert materials["image_requirements"] == [{
        "kind": "reference_acquisition", "mode": "one_shot",
        "purpose": "source_backed_evidence", "request_id": expected_id,
        "material_id": expected_id, "search_query": "growth evidence",
    }]
    assert "Change the revenue fact" not in json.dumps(materials)
    assert "Use attachment attachment-01" not in json.dumps(materials)


def test_long_first_paragraph_is_not_promoted_to_a_fixed_title(tmp_path: Path):
    word = tmp_path / "input.docx"
    logo = tmp_path / "logo.svg"
    project = tmp_path / "project"
    paragraph = "聚焦港澳、内地、国际三地市场，以承建地产为主业，培育实业、金融投资业务，实现利润10%增长。"
    document = Document()
    document.add_paragraph("第1页")
    document.add_paragraph(paragraph)
    document.save(word)
    logo.write_text('<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 100 20"/>', encoding="utf-8")

    state = initialize_v6_project(word, logo, project)
    effective = json.loads((project / "02_v6/effective_pages/page_001.json").read_text(encoding="utf-8"))

    assert state["pages"][0]["title"] != paragraph
    assert len(state["pages"][0]["title"]) <= 28
    assert effective["body_render_content"] == paragraph
    assert effective["word_original"] == paragraph


def test_initialize_v6_project_preserves_embedded_image_integrity_and_paths(tmp_path: Path):
    word = tmp_path / "input.docx"
    logo = tmp_path / "logo.svg"
    project = tmp_path / "project"
    image_path = tmp_path / "source.bmp"
    Image.new("RGB", (8, 6), color=(20, 40, 60)).save(image_path)
    document = Document()
    document.add_paragraph("第1页")
    document.add_paragraph("图片页")
    document.add_paragraph("正文")
    document.add_picture(str(image_path))
    document.save(word)
    logo.write_text('<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 100 20"/>', encoding="utf-8")

    initialize_v6_project(word, logo, project)

    assets = json.loads((project / "02_v6" / "source_assets.json").read_text(encoding="utf-8"))
    asset = next(item for item in assets["assets"] if item["media_type"] == "image/bmp")
    material = json.loads(
        (project / "02_v6" / "page_materials" / "page_001.json").read_text(encoding="utf-8")
    )
    reference = material["reference_images"][0]
    generation_input = asset["generation_input"]

    assert reference["original_path"].startswith("02_v6/reference_media/word_asset_001/original.")
    assert reference["model_input_path"].startswith("02_v6/reference_media/word_asset_001/model-input.")
    assert reference["original_path"] != reference["model_input_path"]
    assert reference["integrity"]["original_sha256"] == asset["sha256"]
    assert reference["integrity"]["model_input_sha256"] != generation_input["sha256"]
    assert reference["thumbnail_path"].endswith("thumbnail.png")
    assert len(reference["integrity"]["thumbnail_sha256"]) == 64


def test_initialize_v6_project_wires_source_chart_records_as_text_facts_only(tmp_path: Path, monkeypatch):
    """Leaving extracted charts outside page materials would make chart_to_facts dead code."""
    word = tmp_path / "input.docx"
    logo = tmp_path / "logo.svg"
    project = tmp_path / "project"
    document = Document()
    document.add_paragraph("Page 1")
    document.add_paragraph("Chart page")
    document.add_paragraph("Narrative")
    document.save(word)
    logo.write_text('<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 100 20"/>', encoding="utf-8")

    def extract_with_chart(_word, _pages, _output):
        return {"assets": [], "chart_records": [{
            "page_numbers": [1], "title": "Revenue trend", "unit": "USD m",
            "series": [{"series": "Revenue", "time": "2025", "value": 20, "trend": "up"}],
            "image_path": "must-not-persist.png",
        }]}

    monkeypatch.setattr(workflow_v6_source, "extract_source_assets", extract_with_chart)
    monkeypatch.setattr(workflow_v6_source, "extract_auto", lambda *_args, **_kwargs: {
        "pagination_mode": "marker", "pages": [{
            "page_number": 1,
            "blocks": [
                {"type": "paragraph", "text": "Chart page", "source_block_index": 0},
                {"type": "paragraph", "text": "Narrative", "source_block_index": 1},
            ],
            "page_comments": [],
        }],
    })
    initialize_v6_project(word, logo, project)

    materials = json.loads((project / "02_v6/page_materials/page_001.json").read_text(encoding="utf-8"))
    assert materials["chart_facts"] == [{
        "title": "Revenue trend", "unit": "USD m",
        "series": [{"series": "Revenue", "time": "2025", "value": 20, "trend": "up"}],
    }]
    assert materials["reference_images"] == []


def test_initialize_v6_project_uses_text_derivative_for_binary_attachment(tmp_path: Path, monkeypatch):
    """Decoding an XLSX original as text would discard available extracted evidence."""
    word = tmp_path / "input.docx"
    logo = tmp_path / "logo.svg"
    project = tmp_path / "project"
    document = Document()
    document.add_paragraph("Page 1")
    paragraph = document.add_paragraph("Attachment evidence")
    document.add_comment(
        [paragraph.runs[0]], "Use attachment workbook rows 2 fields Revenue.",
        author="Reviewer", initials="RV",
    )
    document.save(word)
    logo.write_text('<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 100 20"/>', encoding="utf-8")

    def extract_with_workbook(_word, _pages, output):
        original = output / "00_source/word_assets/original/workbook.xlsx"
        derivative = output / "00_source/word_assets/derived/workbook.txt"
        original.parent.mkdir(parents=True, exist_ok=True)
        derivative.parent.mkdir(parents=True, exist_ok=True)
        original.write_bytes(b"PK binary workbook")
        derivative.write_text("header\nRevenue=20\n", encoding="utf-8")
        return {"assets": [{
            "asset_id": "workbook", "page_numbers": [1],
            "media_type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "relative_path": "00_source/word_assets/original/workbook.xlsx", "sha256": "a" * 64,
            "generation_input": {
                "relative_path": "00_source/word_assets/derived/workbook.txt", "sha256": "b" * 64,
                "media_type": "text/plain", "derivation": "text_extraction",
            },
        }]}

    monkeypatch.setattr(workflow_v6_source, "extract_source_assets", extract_with_workbook)
    monkeypatch.setattr(workflow_v6_source, "extract_auto", lambda *_args, **_kwargs: {
        "pagination_mode": "marker", "pages": [{
            "page_number": 1,
            "blocks": [
                {"type": "paragraph", "text": "Attachment evidence", "source_block_index": 0},
            ],
            "page_comments": [{
                "comment_id": "rows", "text": "Use attachment workbook rows 2 fields Revenue.",
            }],
        }],
    })
    initialize_v6_project(word, logo, project)

    materials = json.loads((project / "02_v6/page_materials/page_001.json").read_text(encoding="utf-8"))
    extracted = materials["attachment_extracts"][0]
    assert extracted["status"] == "available"
    assert extracted["content"] == ["Revenue=20"]
    assert extracted["source_identity"] == {
        "original_path": "01_source_assets/00_source/word_assets/original/workbook.xlsx",
        "original_sha256": "a" * 64,
    }
    assert not any(item["code"] == "attachment_unavailable" for item in materials["degradations"])
