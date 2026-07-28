"""Regression tests for the current-only paginated Word prepare contract."""

from __future__ import annotations

import inspect
import json
import subprocess
import sys
from pathlib import Path

import pytest
from docx import Document
from jsonschema import Draft202012Validator
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
sys.path.insert(0, str(SCRIPTS))

import prepare_run  # noqa: E402
from extract_docx_pages import extract_auto  # noqa: E402
from page_assets import classify_page_asset  # noqa: E402


def make_marked_word(path: Path) -> Path:
    document = Document()
    document.add_paragraph("第1页")
    document.add_paragraph("第一页内容必须留在第一页。")
    document.add_paragraph("第2页")
    document.add_paragraph("第二页内容必须留在第二页。")
    document.save(path)
    return path


def test_prepare_accepts_only_word_and_output_and_locks_ordered_markers(tmp_path: Path):
    """Removing either public argument or reordering marker pages is a contract break."""
    word = make_marked_word(tmp_path / "source.docx")
    project = tmp_path / "project"

    assert list(inspect.signature(prepare_run.prepare).parameters) == ["word", "output"]
    result = prepare_run.prepare(word, project)

    pages = json.loads((project / "00_source" / "pages.json").read_text(encoding="utf-8"))
    assert result["page_count"] == 2
    assert pages["pagination_mode"] == "explicit_text_markers"
    assert [page["page_number"] for page in pages["pages"]] == [1, 2]
    assert pages["pages"][0]["blocks"][0]["text"] == "第一页内容必须留在第一页。"
    assert pages["pages"][1]["blocks"][0]["text"] == "第二页内容必须留在第二页。"


def test_prepare_cli_creates_a_project_without_legacy_inputs(tmp_path: Path):
    """Requiring or silently accepting a legacy input breaks the public CLI contract."""
    word = make_marked_word(tmp_path / "source.docx")
    project = tmp_path / "project"
    command = [
        sys.executable,
        str(SCRIPTS / "word_to_editable_ppt.py"),
        "prepare",
        "--word",
        str(word),
        "--output",
        str(project),
    ]

    completed = subprocess.run(command, capture_output=True, text=True, check=False)

    assert completed.returncode == 0, completed.stderr
    assert json.loads(completed.stdout)["page_count"] == 2
    assert (project / "workflow_run.json").is_file()

    rejected = subprocess.run(
        [*command, "--style-reference", "unused.png", "--company-logo", "unused.png"],
        capture_output=True,
        text=True,
        check=False,
    )
    assert rejected.returncode != 0
    assert "unrecognized arguments" in rejected.stderr


def test_workflow_cli_waits_for_style_confirmation_and_rejects_legacy_commands(tmp_path: Path):
    """A fresh project must not enter a removed master, sample, or global-QA workflow."""
    project = tmp_path / "project"
    prepare_run.prepare(make_marked_word(tmp_path / "source.docx"), project)
    command = [sys.executable, str(SCRIPTS / "word_to_editable_ppt.py"), "workflow"]

    for workflow_command in ["next", "status", "resume"]:
        completed = subprocess.run(
            [*command, workflow_command, "--project", str(project)],
            capture_output=True,
            text=True,
            check=False,
        )
        assert completed.returncode == 0, completed.stderr
        assert json.loads(completed.stdout)["stage"] == "await_style_confirmation"

    for legacy_command in ["confirm-master", "record-sample", "record-global-qa"]:
        rejected = subprocess.run(
            [*command, legacy_command, "--project", str(project)],
            capture_output=True,
            text=True,
            check=False,
        )
        assert rejected.returncode != 0
        assert "invalid choice" in rejected.stderr


def test_physical_fallback_is_used_only_when_markers_are_absent(tmp_path: Path, monkeypatch: pytest.MonkeyPatch):
    """Invalid markers must remain an error; only an unmarked document may use physical pagination."""
    unmarked = Document()
    unmarked.add_paragraph("没有页码标记的正文")
    unmarked_path = tmp_path / "unmarked.docx"
    unmarked.save(unmarked_path)
    physical_pages = {
        "schema_version": "1.1",
        "source_file": unmarked_path.name,
        "pagination_mode": "physical_rendered_pages",
        "pagination_backend": "test-renderer",
        "page_count": 1,
        "pages": [{"page_number": 1, "blocks": [{"type": "paragraph", "text": "物理分页内容"}], "must_keep": [], "page_purpose": "待人工填写"}],
    }
    monkeypatch.setattr("extract_docx_pages.extract_physical", lambda _path: physical_pages)

    assert extract_auto(unmarked_path) == physical_pages

    invalid = Document()
    invalid.add_paragraph("第2页")
    invalid.add_paragraph("顺序错误")
    invalid_path = tmp_path / "invalid.docx"
    invalid.save(invalid_path)
    with pytest.raises(ValueError, match="ordered and consecutive"):
        extract_auto(invalid_path)


def test_prepared_state_has_the_word_only_shape(tmp_path: Path):
    """Reintroducing old master, sample, logo, or visual-DNA state is a contract break."""
    project = tmp_path / "project"
    prepare_run.prepare(make_marked_word(tmp_path / "source.docx"), project)

    state = json.loads((project / "workflow_run.json").read_text(encoding="utf-8"))
    assert {"workflow_contract_version", "word_source", "pagination", "style_confirmation", "jobs", "final_pptx"} <= state.keys()
    assert not {"master_jobs", "sample_status", "company_logo", "visual_dna_receipt"} & state.keys()
    assert state["word_source"]["path"] == "00_source/source.docx"
    assert state["pagination"]["locked_page_order"] == [1, 2]
    assert state["style_confirmation"]["status"] == "pending"
    assert [job["page_number"] for job in state["jobs"]] == [1, 2]
    assert [job["expected_output"] for job in state["jobs"]] == [
        "06_images/generated/page_001.png",
        "06_images/generated/page_002.png",
    ]
    assert state["final_pptx"] is None

    template = ROOT / "template"
    assert (template / "06_images" / "generated").is_dir()
    assert (template / "08_final").is_dir()
    assert not (template / "06_images" / "approved").exists()
    assert not (template / "06_images" / "draft").exists()
    assert not (template / "09_deliverables").exists()


def test_inline_word_image_is_extracted_and_bound_only_to_its_locked_page(tmp_path: Path):
    """A Word image must never leak into another independently generated slide."""
    image_path = tmp_path / "chart.png"
    Image.new("RGB", (80, 40), "#22577A").save(image_path)
    document = Document()
    document.add_paragraph("第1页")
    document.add_paragraph("根据下图说明第一阶段进展。")
    document.add_picture(str(image_path))
    document.add_paragraph("第2页")
    document.add_paragraph("第二页不应收到第一页图片。")
    word = tmp_path / "with-image.docx"
    document.save(word)

    project = tmp_path / "project"
    prepare_run.prepare(word, project)
    first = json.loads((project / "01_page_contracts/page_001.json").read_text(encoding="utf-8"))
    second = json.loads((project / "01_page_contracts/page_002.json").read_text(encoding="utf-8"))

    assert len(first["asset_bindings"]) == 1
    binding = first["asset_bindings"][0]
    assert binding["asset_role"] == "visual_reference"
    assert binding["processing"] == "direct_image"
    assert binding["use_policy"] == "required"
    assert binding["advisories"] == []
    assert (project / binding["generation_input"]["relative_path"]).read_bytes() == image_path.read_bytes()
    assert second["asset_bindings"] == []
    manifest = json.loads((project / "00_source/source_asset_manifest.json").read_text(encoding="utf-8"))
    schema = json.loads((ROOT / "schemas/source_asset_manifest.schema.json").read_text(encoding="utf-8"))
    assert not list(Draft202012Validator(schema).iter_errors(manifest))


@pytest.mark.parametrize(
    ("media_type", "expected_role", "expected_processing"),
    [
        ("application/pdf", "document_source", "extract_content"),
        ("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "data_source", "extract_content"),
        ("application/vnd.openxmlformats-officedocument.presentationml.presentation", "document_source", "extract_content"),
        ("application/octet-stream", "unsupported", "unavailable"),
    ],
)
def test_page_attachment_classification_is_explicit_and_non_blocking(
    media_type: str, expected_role: str, expected_processing: str
):
    """Unreadable attachments should be visible to QA without stopping other pages."""
    result = classify_page_asset(media_type, binding_status="bound", has_generation_input=False)
    assert result["asset_role"] == expected_role
    assert result["processing"] == expected_processing
    assert result["blocking"] is False
    assert bool(result["advisories"]) is (expected_processing == "unavailable")
